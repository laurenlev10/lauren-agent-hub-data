"""
meta_inbox_preview — Phase 1 dry-run for the @meta agent migration.

Fetches Lauren's actual inbox via the Meta Graph API (Messenger DMs + FB/IG
comments), classifies each item against the KB, and generates a static HTML
preview at docs/meta/inbox-api-preview/index.html showing what the API-based
agent WOULD reply if it ran in production.

Lauren reviews the preview, compares classifications to her gut feel, and
either approves the cutover or flags items that look wrong.

This is a *read-only* run — no replies are sent, no comments are hidden.
"""

import datetime as dt
import html
import json
import os
import re
import sys
from pathlib import Path

# Add scripts/ to path so we can import lauren_meta
sys.path.insert(0, str(Path(__file__).parent))
from lauren_meta import fetch_recent_inbox, fetch_messenger_messages, get_token, get_fb_page_id

from docx import Document


# --- KB loader ---------------------------------------------------------------


# ============================================================
# 2026-05-14 PM — Sentiment + Urgency classifiers (keyword-based)
# Lauren's directive: tag every incoming message, SMS on urgent.
# ============================================================

_SENTIMENT_KEYWORDS = {
    "angry": [
        # Hebrew
        "כועס", "מאוכזב", "תלונה", "תרמיתם", "רמאות", "גנבים", "גרוע", "נורא",
        "אסון", "מטופש", "לא מקצועי", "חוצפה", "בושה", "בעיה",
        # English
        "angry", "furious", "scam", "fraud", "rip off", "ripoff", "terrible",
        "awful", "horrible", "complaint", "disappointed", "refund", "lawsuit",
        "report", "bbb", "stolen", "stole",
    ],
    "complaint": [
        # Hebrew
        "תלונה", "החזר", "כסף בחזרה", "לא מרוצה", "לא מקצועי", "בעיה",
        # English
        "complaint", "refund", "return", "cancel", "unhappy", "not happy",
        "issue", "problem", "broken", "defective", "damaged",
    ],
    "positive": [
        # Hebrew
        "תודה", "מעולה", "אהבתי", "מדהים", "מקסים", "מומלץ", "שיא", "אש",
        # English
        "thank you", "thanks", "thx", "love", "amazing", "perfect", "awesome",
        "great", "wonderful", "recommend", "fire", "best",
    ],
}

_URGENT_KEYWORDS = [
    # Hebrew
    "דחוף", "מיידי", "עכשיו", "תלונה", "החזר", "תרמיתם", "כועס", "מאוכזב",
    "מתי תתפתחו", "מתי הם פותחים", "סגור", "בעיה דחופה",
    # English
    "urgent", "asap", "immediately", "complaint", "refund", "scam", "fraud",
    "stolen", "bbb", "lawsuit", "police", "report",
]


def classify_sentiment(text: str) -> str:
    """Return one of: angry | complaint | positive | neutral."""
    if not text:
        return "neutral"
    t = text.lower()
    # Priority: angry > complaint > positive > neutral
    for kw in _SENTIMENT_KEYWORDS["angry"]:
        if kw in t:
            return "angry"
    for kw in _SENTIMENT_KEYWORDS["complaint"]:
        if kw in t:
            return "complaint"
    for kw in _SENTIMENT_KEYWORDS["positive"]:
        if kw in t:
            return "positive"
    return "neutral"


def is_urgent(text: str) -> bool:
    """Return True if message contains urgency-signaling keywords."""
    if not text:
        return False
    t = text.lower()
    return any(kw in t for kw in _URGENT_KEYWORDS)


def load_venues() -> list:
    """Parse scripts/data/venue_details.md into list of {dates, city, venue, address} dicts."""
    candidates = [
        Path(__file__).resolve().parent / "data" / "venue_details.md",
        Path("/sessions/dreamy-compassionate-wozniak/mnt/Claude/Scheduled/NEW/meta-inbox/venue_details.md"),
    ]
    for c in candidates:
        if not c.exists():
            continue
        out = []
        for line in c.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if not line.startswith("|") or "---" in line or line.startswith("| Event") or line.startswith("| dates"):
                continue
            cells = [x.strip() for x in line.split("|")[1:-1]]
            if len(cells) == 4 and cells[0]:
                out.append({"dates": cells[0], "city": cells[1], "venue": cells[2], "address": cells[3]})
        return out
    return []


def find_event_today(venues, today=None):
    today = today or dt.date.today()
    for v in venues:
        start, end, _ = _parse_dates(v["dates"])
        if start and end and start <= today <= end:
            return v
    return None


def find_next_event(venues, today=None):
    today = today or dt.date.today()
    upcoming = [(_parse_dates(v["dates"])[0], v) for v in venues if _parse_dates(v["dates"])[0] and _parse_dates(v["dates"])[0] > today]
    if not upcoming:
        return None
    upcoming.sort(key=lambda x: x[0])
    return upcoming[0][1]


def load_handled() -> dict:
    """
    Load docs/meta/handled.json — the per-item dedup memory.
    Schema: { "<channel>:<id>": {"handled": True, "handledAt": "ISO"} }
    Updated by the agent (after auto-reply) AND by Lauren's clicks in the preview UI
    (which calls GitHub Contents API to PUT the file).
    """
    candidates = [
        Path(__file__).resolve().parent.parent / "docs/meta/handled.json",
    ]
    for c in candidates:
        if c.exists():
            try:
                return json.loads(c.read_text(encoding="utf-8"))
            except Exception:
                return {}
    return {}


def dedup_key(channel: str, item_id: str) -> str:
    """Stable key for an inbox item across runs."""
    return f"{channel}:{item_id}"


def is_handled(handled: dict, channel: str, item_id: str) -> bool:
    """True if this exact item has been marked handled in a prior run."""
    if not handled or not item_id:
        return False
    entry = handled.get(dedup_key(channel, item_id), {})
    return bool(entry.get("handled"))


def load_kb(docx_path: Path) -> dict:
    """Parse the inbox_knowledge_base.docx into a structured KB."""
    doc = Document(str(docx_path))
    kb = {"faqs": [], "schedule": {}, "negatives": []}

    # The docx has 4 tables. Order is fixed:
    # 0: FAQ (Question, Standard Reply)
    # 1: Event schedule (Dates, City)
    # 2: Negative/skeptical situations (Situation, How to respond)
    # 3: Optional fourth (links/etc)

    if doc.tables:
        # FAQ table
        for row in doc.tables[0].rows[1:]:  # skip header
            q = row.cells[0].text.strip()
            a = row.cells[1].text.strip()
            if q and a:
                kb["faqs"].append({"q": q, "a": a})

    if len(doc.tables) >= 2:
        for row in doc.tables[1].rows[1:]:
            dates = row.cells[0].text.strip()
            city = row.cells[1].text.strip()
            if dates and city:
                kb["schedule"][city.lower()] = dates

    if len(doc.tables) >= 3:
        for row in doc.tables[2].rows[1:]:
            sit = row.cells[0].text.strip()
            resp = row.cells[1].text.strip()
            if sit and resp:
                kb["negatives"].append({"situation": sit, "reply": resp})

    return kb


# --- Classifier --------------------------------------------------------------

# Keyword → which FAQ to match. Tuned for Lauren's actual common questions.
FAQ_KEYWORDS = [
    (re.compile(r"\b(free|admission|ticket|cost.*entry|enter)\b", re.I),
     "Is admission free?"),
    (re.compile(r"\b(what\s+(time|hour)|when.*(open|close|hour)|"
                r"(hours?|times?)\s+of\s+operation|"
                r"\d{1,2}\s*(am|pm)|10\s*-\s*5)\b", re.I),
     "What are the hours?"),
    (re.compile(r"\b(brand|carry|sell|product line)s?\b", re.I),
     "What brands do you have?"),
    (re.compile(r"\b(what.*mystery\s*box|mystery\s*box.*what|how.*mystery\s*box|mystery\s*box\s*\?|grab\s*bag)\b", re.I),
     "What is the Mystery Box?"),
    (re.compile(r"\b(parking|lot)\b", re.I),
     None),  # FAQ may have "Is parking free?"; if not matched will be Bucket B
    (re.compile(r"\b(payment|cash|card|credit|debit|tap)\b", re.I),
     None),
    (re.compile(r"\b(kids|child|stroller|baby)\b", re.I),
     None),
    (re.compile(r"\b(refund|return)s?\b", re.I),
     None),
]

NEGATIVE_KEYWORDS = re.compile(
    r"\b(scam|fake|fraud|rip[\s-]?off|too good to be true|bots?|spam)\b",
    re.I,
)

# Customer-issue patterns — past purchase / problem report. ALWAYS Bucket B.
CUSTOMER_ISSUE_PAT = re.compile(
    r"\b(i\s+(bought|purchased|got|paid|ordered)|"
    r"problem|issue|complaint|didn'?t\s+(receive|get|work)|"
    r"never\s+(received|got)|broken|damaged|wrong\s+item|"
    r"refund|return|missing|disappointed|unhappy)\b",
    re.I,
)

# "Can I order online?" / "Available online?" — direct her to the website
# "Where?" comments on event-specific reels — answer with the venue from the post caption.
WHERE_PAT = re.compile(
    r"^\s*(?:at\s+)?where\??!?\s*\??!?\s*$|"        # "Where?" / "At where??"
    r"^\s*where\s+(?:is|are|r|u|you|it|this|the\s+sale)\b",  # "where is it"
    re.IGNORECASE,
)

# Patterns that pull event info out of Lauren's reel captions
_CITY_STATE_FROM_CAP = re.compile(
    r"sale\s+in\s+([^,!?\n]+?)\s*,\s*([A-Z]{2})", re.IGNORECASE
)
_WHEN_FROM_CAP = re.compile(r"when[:\s]+([^\n]+?)\s*\n", re.IGNORECASE)
_WHERE_FROM_CAP = re.compile(r"where[:\s]+([^\n]+?)\s*\n", re.IGNORECASE)
_AT_VENUE_FROM_CAP = re.compile(r"\bAt\s+([A-Z][^.\n]+?)(?:\.|\n|$)")


def summarize_event_from_caption(caption: str):
    """Pull (city, state, dates, address, venue) out of one of Lauren's
    standard event reel captions. Returns None if we can't identify a city."""
    if not caption:
        return None
    m = _CITY_STATE_FROM_CAP.search(caption)
    if not m:
        return None
    city, state = m.group(1).strip(), m.group(2).strip().upper()
    dates_m   = _WHEN_FROM_CAP.search(caption)
    address_m = _WHERE_FROM_CAP.search(caption)
    venue_m   = _AT_VENUE_FROM_CAP.search(caption)
    return {
        "city":    city,
        "state":   state,
        "dates":   dates_m.group(1).strip() if dates_m else None,
        "address": address_m.group(1).strip() if address_m else None,
        "venue":   venue_m.group(1).strip() if venue_m else None,
    }


def _enrich_event_info(info: dict, venues: list) -> dict:
    """If `info` from summarize_event_from_caption has city/state but lacks
    dates/address/venue (common when post caption is short), look up the
    canonical schedule entry by city match and fill in the missing fields.

    Returns the enriched dict (mutated). 2026-05-20 PM #2.
    """
    if not info or not info.get("city") or not venues:
        return info
    city_l = info["city"].strip().lower()
    state_l = (info.get("state") or "").strip().upper()
    for v in venues:
        v_city_full = (v.get("city") or "").strip()
        # "Overland Park, KS" → "overland park", "KS"
        v_city = v_city_full.split(",")[0].strip().lower()
        v_state = v_city_full.split(",")[1].strip().upper() if "," in v_city_full else ""
        if v_city == city_l or city_l in v_city or v_city in city_l:
            if state_l and v_state and state_l != v_state:
                continue  # state mismatch (e.g., Roseville MN vs Roseville CA)
            if not info.get("dates"):   info["dates"]   = v.get("dates")
            if not info.get("address"): info["address"] = v.get("address")
            if not info.get("venue"):   info["venue"]   = v.get("venue")
            break
    return info


def event_status(dates_str: str, today: dt.date = None) -> str:
    """Returns 'live' / 'upcoming-soon' / 'upcoming' / 'past' / 'unknown'."""
    today = today or dt.date.today()
    start, end, _ = _parse_dates(dates_str or "")
    if not start or not end:
        return "unknown"
    if today < start:
        days = (start - today).days
        return "upcoming-soon" if days <= 7 else "upcoming"
    if today > end:
        return "past"
    return "live"  # today between start and end inclusive


_WHERE_LIVE_TEMPLATES = [
    "We're LIVE NOW in {city}, {state}!! 💄✨ 📍 {address} — open today "
    "10am-5pm. Come thru gorgeous 💕",
    "Happening RIGHT NOW 💄 {city}, {state} — 📍 {address}, open till 5pm. "
    "Don't sleep on this 🛍️✨",
    "Babe we're LIVE 🎉 {city}, {state} — 📍 {address} (at {venue}). "
    "10am-5pm today, free entry + parking 💄💕",
    "Today's the day!! 🥳 {city}, {state} — 📍 {address}. Open 10am-5pm 💄✨",
]
_WHERE_UPCOMING_SOON_TEMPLATES = [
    "Coming up THIS weekend 🎉 {city}, {state} — 📍 {address}, Fri-Sun "
    "10am-5pm 💄✨ Free entry + parking 💕",
    "We're heading to {city}, {state} — {dates}!! 📍 {address}. "
    "See you there gorgeous 💄✨",
    "Locked in for {city}, {state} — {dates} 💄 📍 {address}. "
    "Fri-Sun 10am-5pm, free admission 💕",
]
_WHERE_UPCOMING_TEMPLATES = [
    "{city}, {state} — {dates} 💄 📍 {address}. Fri-Sun 10am-5pm. "
    "Mark your calendar 📌💕",
    "We're popping up in {city}, {state} for {dates} 🎉 📍 {address}. "
    "Free entry, free parking 💄✨",
    "Heading to {city}, {state} — {dates}! 📍 {address} (at {venue}). "
    "Fri-Sun 10am-5pm 💄✨",
]
OPEN_TODAY_PAT = re.compile(
    r"\b(?:open\s+today|here\s+today|live\s+today|are\s+you\s+open|"
    r"are\s+you\s+here|you\s+open\s+today|today\s+open)\b",
    re.IGNORECASE,
)
PAYMENT_PAT = re.compile(
    r"\b(?:cash\s+only|take\s+(?:credit|debit|cards?|cash|apple\s*pay|tap)|"
    r"accept\s+(?:credit|debit|cards?|cash|apple\s*pay|tap)|"
    r"do\s+you\s+(?:take|accept)\s+(?:credit|debit|cards?|cash|venmo|zelle)|"
    r"payment\s+method|tap\s*to\s*pay|apple\s*pay)\b",
    re.IGNORECASE,
)
PARKING_PAT = re.compile(r"\b(?:parking|where\s+(?:do|can)\s+i\s+park|valet)\b", re.IGNORECASE)
HOURS_PAT = re.compile(
    r"\b(?:what\s+time|what\s+(?:are\s+the\s+)?hours?|when\s+(?:do\s+you\s+)?(?:open|close)|"
    r"hours?\s+of\s+operation|opening\s+hours?)\b",
    re.IGNORECASE,
)
KIDS_PAT = re.compile(
    r"\b(?:bring\s+(?:my\s+)?(?:kids?|children|baby|toddler|stroller)|"
    r"kid[s]?\s+(?:ok|allowed|welcome|friendly)|kid[\s-]?friendly|"
    r"family[\s-]?friendly|stroller)\b",
    re.IGNORECASE,
)
ADDRESS_PAT = re.compile(
    r"\b(?:what'?s?\s+the\s+(?:address|location|venue)|address(?:\s|\?|$)|"
    r"location(?:\s|\?|$)|venue\s*\?|exact\s+(?:address|location)|"
    r"where\s+exactly|directions\??)\b",
    re.IGNORECASE,
)


_OPEN_TODAY_LIVE_TEMPLATES = [
    "YES we're LIVE NOW!! 💄✨ Open till 5pm at {address} ({city}, {state}) — come thru gorgeous 💕",
    "OPEN!! 🎉 We're at {address} ({city}, {state}) until 5pm tonight — bring your bestie 💄✨",
    "Yes ma'am 💄 LIVE today at {address}, open till 5pm. See you soon 💕",
    "🔥 Live + open!! {city}, {state} — 📍 {address}, until 5pm 💄✨",
]
_OPEN_TODAY_CLOSED_TEMPLATES = [
    "Not today girl 💔 But coming up: {city}, {state} — {dates}!! 📍 {address}. See you there 💄✨",
    "Aww not open today 😢 Next stop: {city}, {state} ({dates}) — 📍 {address}. Mark your calendar 📌💄",
    "Closed today, but {city}, {state} is up next ({dates}) — 📍 {address} 💄✨",
]
_PAYMENT_TEMPLATES = [
    "We take EVERYTHING — cash, cards, debit, tap-to-pay, Apple Pay 💳✨ However works for you 💄",
    "All payment methods welcome! 💳 Cash, cards, tap-to-pay, Apple Pay — you name it 💄✨",
    "Cash, cards, tap-to-pay — we accept it all 💳💄 Whatever's easiest 💕",
    "Yes! 💳 Cash, credit, debit, tap-to-pay, Apple Pay — all good 💄✨",
]
_PARKING_TEMPLATES = [
    "Yes!! 🚗 Free parking AND free entry — bring the whole gang 💄✨",
    "Parking is FREE 🅿️ And so is admission — just show up & shop 💄💕",
    "Free parking, free entry, free vibes 💄✨ See you there 💕",
    "🅿️ Yep, parking is on the house — and admission too 💄✨",
]
_HOURS_TEMPLATES = [
    "Friday, Saturday & Sunday — 10am to 5pm ⏰✨ See you there 💄",
    "We're open Fri-Sun 10am-5pm 💄 Plenty of time to come browse 💕",
    "10am-5pm Friday through Sunday ⏰💄 Free entry, no ticket needed ✨",
    "Fri/Sat/Sun, 10am-5pm 💄✨ Drop by anytime in those hours 💕",
]
_KIDS_TEMPLATES = [
    "Yes!! 👶 Kids are welcome — bring the whole crew 💄💕",
    "Of course gorgeous 💄 The whole family is welcome, kids included 👶✨",
    "All ages welcome 💕 Bring the kids, the bestie, your mama 💄✨",
    "Kid-friendly all the way 👶💄 Strollers, the works — come on by ✨",
]


_WHERE_PAST_TEMPLATES = [
    "Aww we already wrapped {city} {dates} 😭💔 — but watch our FB for "
    "the next stop near you! 💄✨",
    "{city} just wrapped {dates} 😭 — we rotate cities every 1–2 years, "
    "but stay close on FB so you catch the next one 💄💕",
]


ONLINE_ORDER_PAT = re.compile(
    r"\b(?:order|buy|purchase|shop|sell|sold|ship|delivery|delivered|"
    r"available|get\s+it)\b.*\bonline\b|"
    r"\bonline\b.*\b(?:order|buy|purchase|shop|sell|sold|ship|delivery|"
    r"delivered|available)\b",
    re.I,
)

_ONLINE_ORDER_TEMPLATES = [
    "YES girl!! 💄 You can grab our Mystery Box online anytime: "
    "https://www.themakeupblowout.com/ ✨💕",
    "Heck yes 🛍️ Online shop right here: https://www.themakeupblowout.com/ — "
    "same fabulous deals delivered to your door 📦✨",
    "OMG yes queen! 💄 Mystery Box online: https://www.themakeupblowout.com/ "
    "✨ Bring the magic home 💕",
    "Lovely yes!! 💕 Online store: https://www.themakeupblowout.com/ — "
    "delivered to your door 📦💄✨",
    "Hey gorgeous 💄 Online shop's right here: https://www.themakeupblowout.com/ "
    "✨ Same magic, delivered 💕",
    "Babe YESSS!! 🛍️ https://www.themakeupblowout.com/ — same fabulous deals "
    "delivered to your door 📦💄💕",
]

CITY_PAT = re.compile(r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b")


# Pattern that detects "when are you coming to X / X event / sale in X"
# style intent, regardless of capitalization.
CITY_QUESTION_PAT = re.compile(
    r"\b(when|are you|coming|visit|going|sale|event|come|been|be at|make it to)\b.*"
    r"\b(?:to|in|at|near|around)\s+"
    # Capture up to ?, sentence boundary, or a clearly terminal word.
    # Allows multi-word places like "Northern California bay area" or
    # "San Bernardino county". Terminates at a new clause boundary
    # (e.g. " I need" / " love your" / " thanks") so we don't swallow
    # the rest of the message.
    r"([A-Za-z][A-Za-z0-9\s.\-,/]{1,80}?)"
    r"(?:\?|\.|!|$|"
    r"\s+I[\s\u2019']|"     # "Palmdale / Lancaster I need..." / "I'm" / "I've"
    r"\s+(?:love|need|want|hope|wish|please|thanks?|tho|though|thx|ty)\b)",
    re.IGNORECASE,
)

# Also catch the bare "X???" pattern — just a place name, no verb (common on mobile).
# We only fire this fallback if no FAQ keyword matched AND the text is short.
BARE_PLACE_PAT = re.compile(r"^\s*([A-Za-z][\w\s.\-,/]{2,40}?)\s*\?+\s*$", re.IGNORECASE)

# 2026-05-14 PM — relaxed place mention without question word.
# Matches "(any prefix)come/back/return/visit/miss(any) to/in/at <Place>"
RELAXED_PLACE_PAT = re.compile(
    r"\b(?:come|back|return|visit|miss|love)\b[^A-Za-z]*?"
    r"\b(?:to|in|at|out\s+in)\s+"
    r"([A-Z][\w][\w\s.\-,/]{1,40}?)"
    # 2026-05-20 — added [^\w\s.\-,/] so emoji/symbol terminates (was failing on "Victorville🥲")
    r"(?:[\?\.\!,]|\b(?:soon|please|again|next|this|that|love|need|want|hope|thanks?)\b|[^\w\s.\-,/]|$)",
    re.IGNORECASE,
)

# 2026-05-20 PM #2 — "Interested in this event" / "I'll be there" / "can't wait"
# When this matches on a post with event context, draft an event-specific reply
# with venue + dates + cute closing (per Lauren's directive — see memory.md).
INTERESTED_PAT = re.compile(
    r"\b(?:interested\s+(?:in|on)|i'?m\s+(?:interested|in|coming|going|there)|"
    r"can'?t\s+wait|i'?ll\s+(?:be\s+there|come|attend|see\s+you)|"
    r"i\s+(?:want|need|wanna)\s+(?:to\s+)?(?:come|go|attend|be\s+there)|"
    r"sign\s+me\s+up|count\s+me\s+in|see\s+you\s+there|i'?m\s+going|"
    r"this\s+is\s+(?:so\s+)?(?:cool|cute|exciting|amazing)|"
    r"omg\s+(?:yes|yas+|i\s+need)|need\s+this|love\s+this\s+event)\b",
    re.IGNORECASE,
)



def _normalize_place(s: str) -> str:
    """Lower, strip punctuation, collapse spaces."""
    return re.sub(r"[^a-z0-9 ]+", " ", s.lower()).strip()


# Common US-place prefixes that often get concatenated in mobile typing.
# 2026-07-02 #2 — removed "st " ("stockton" was prettified to "St Ockton") and
# "el " ("elmira" → "El Mira" risk). Short ambiguous prefixes do more harm than good.
_PLACE_PREFIXES = ["san ", "los ", "las ", "new ", "north ", "south ",
                   "east ", "west ", "fort ", "saint "]

def _prettify_place(raw: str) -> str:
    """
    Turn user-typed place strings into a clean Title Case for the reply.
    Examples:
      'sanbernardino'           -> 'San Bernardino'
      'Northern California bay' -> 'Northern California Bay'
      'arizona'                 -> 'Arizona'
    """
    p = raw.strip().rstrip("?,.!").strip()
    pl = p.lower()
    if " " not in pl and pl == p.lower():
        for prefix in _PLACE_PREFIXES:
            compact = prefix.replace(" ", "")
            if pl.startswith(compact) and len(pl) >= len(compact) + 4 and not pl.startswith(prefix):
                rest = p[len(compact):]
                return (prefix + rest).title()
    return p.title()


# --- Date parsing for the schedule -------------------------------------------

_MONTHS = {"jan":1,"january":1,"feb":2,"february":2,"mar":3,"march":3,
           "apr":4,"april":4,"may":5,"jun":6,"june":6,"jul":7,"july":7,
           "aug":8,"august":8,"sep":9,"sept":9,"september":9,"oct":10,"october":10,
           "nov":11,"november":11,"dec":12,"december":12}

def _parse_dates(date_str: str, year: int = 2026):
    """
    Parse a schedule string like 'May 1–3', 'March 13–15', 'Jul 31–Aug 2'
    into (start_date, end_date, end_month_name).
    Returns (None, None, None) if parsing fails.
    """
    if not date_str:
        return None, None, None
    # Normalize dashes
    s = date_str.replace("–", "-").replace("—", "-").strip()
    # Form A: "May 1-3"
    m = re.match(r"^([A-Za-z]+)\s+(\d{1,2})\s*-\s*(\d{1,2})$", s)
    if m:
        mon = _MONTHS.get(m.group(1).lower())
        if mon:
            try:
                d1 = dt.date(year, mon, int(m.group(2)))
                d2 = dt.date(year, mon, int(m.group(3)))
                return d1, d2, m.group(1).title()
            except ValueError:
                pass
    # Form B: "Jul 31-Aug 2"
    m = re.match(r"^([A-Za-z]+)\s+(\d{1,2})\s*-\s*([A-Za-z]+)\s+(\d{1,2})$", s)
    if m:
        mon1 = _MONTHS.get(m.group(1).lower())
        mon2 = _MONTHS.get(m.group(3).lower())
        if mon1 and mon2:
            try:
                d1 = dt.date(year, mon1, int(m.group(2)))
                d2 = dt.date(year, mon2, int(m.group(4)))
                return d1, d2, m.group(3).title()
            except ValueError:
                pass
    return None, None, None


def _is_past(date_str: str, today: dt.date = None) -> bool:
    """True if the event end date is strictly before today."""
    today = today or dt.date.today()
    _, end, _ = _parse_dates(date_str)
    if end is None:
        return False
    return end < today


# --- Reply template pools (seeded by conv_id for consistency) ----------------

import hashlib as _hashlib

_ON_SCHEDULE_TEMPLATES = [
    "YESSS girl!! 🎉 We're rolling into {C} {D} — Friday–Sunday, 10am–5pm. "
    "Free entry, free parking, and a whole lotta beauty waiting for you 💄✨",
    "OMG YES queen 👑 We're hitting {C} {D}! Bring your bestie, your wallet, "
    "and your wishlist — Fri–Sun 10am–5pm, free entry + parking 💄💥",
    "Heck yes mama!! 🥳 {C} is locked in for {D} — Fri/Sat/Sun 10am–5pm. "
    "Free entry, free parking, 75% off your fave brands 💄💥 Mark that calendar 📌",
    "Stop everything 🛑 {C} sale is happening {D}! Fri–Sun 10am–5pm. "
    "Free entry, free parking, the biggest beauty deals you've ever seen 💄✨",
    "Hey gorgeous 💄 {C} sale is on for {D}! Fri–Sun 10am–5pm, free entry. "
    "Show up, no ticket needed, just a love for makeup ✨💕",
    "Lovely yesss!! 💕 We're swinging into {C} {D} — Fri–Sun 10am–5pm. "
    "Free admission, free parking, treat yourself 💄✨",
    "Honey YES!! 🥳 {C} {D} — Fri–Sun 10am–5pm. Free entry + parking, "
    "deals deeper than the ocean 💄💥",
    "Aww babe!! 🎉 Locked in — {C} {D}, Fri–Sun 10am–5pm. Free entry, "
    "free parking. Can't wait to spoil you 💕",
    "Sis YESSS 💄 We're popping up in {C} {D}! Fri–Sun 10am–5pm. "
    "Free admission, no ticket needed — just bring your love for beauty ✨",
    "Doll, locked in!! 👑 {C} sale {D} — Fri–Sun 10am–5pm. Free entry, "
    "free parking, brace yourself for the biggest deals 💄💕",
]

_OFF_SCHEDULE_TEMPLATES = [
    "Aww girl 💔 {C} didn't make our 2026 lineup this year — we visit each city "
    "every 1–2 years to keep things juicy 🤩 Adding {C} to our 2027 wishlist "
    "now 📝💄💕 Follow our Facebook for the announcement 👀",
    "Oof — no {C} on the 2026 schedule this year 😢 We rotate cities to keep "
    "the excitement up (and the deals deeper 💄✨) — but {C} is officially "
    "on our wishlist for 2027 💕 Keep an eye on FB for updates 👀",
    "Sis, {C} isn't on our 2026 stop list this year 💔 We rotate cities every "
    "1–2 years to keep the magic alive ✨ Putting {C} on the radar for 2027 💄💕",
    "Hey love 💕 {C} isn't on the 2026 lineup — but stay close, we'll do our "
    "BEST to come for 2027 💄✨ Watch our Facebook for the announcement 👀",
    "Aww mama 💔 No {C} in 2026 — we rotate cities every 1–2 years to give "
    "each one our ALL when we come ✨ Definitely on our radar for 2027 💄💕",
    "Lovely girl 💕 {C} isn't on our 2026 lineup this year — but trust me, "
    "you're heard. We'll do our best to bring the sale to {C} in 2027 💄✨ "
    "Follow our Facebook for the drop 👀",
    "Gorgeous, {C} didn't make 2026 this year 💔 We rotate cities to keep the "
    "magic alive ✨ But we're putting {C} on the radar HARD for 2027 💄💕 "
    "Keep an eye on FB so you don't miss the announcement 👀",
    "Aww babe — {C} isn't on our 2026 stop list 💔 We visit each city every "
    "1–2 years. {C} is officially on the 2027 wishlist now 📝💄✨ Follow our "
    "Facebook for the announcement 👀",
]

_PAST_SCHEDULE_TEMPLATES = [
    "Aww girl, we were just in {C} {D}! 😭💔 We rotate cities every 1–2 years "
    "so we won't be back in 2026, but you're locked into our 2027 wishlist 💄💕 "
    "Follow our FB so you don't miss the next one ✨",
    "Oof — you JUST missed us!! 😩 {C} happened {D}. We rotate cities so won't "
    "be back this year, but {C}'s on the 2027 radar 💄✨ Stay close on Facebook 👀",
    "Awww mama, we were in {C} {D}! 😭💔 The sale already wrapped — we'll do "
    "our BEST to come back in 2027 💄💕 Watch our FB for the announcement 👀",
    "Hey gorgeous 💕 We were in {C} {D} — already wrapped! We rotate cities, "
    "so won't be back in 2026, but {C}'s officially on the 2027 list 💄✨",
    "Sweetheart 💔 {C} just wrapped {D} — you're officially on our 2027 "
    "wishlist 💄💕 Stay close on FB for the next announcement 👀",
    "Aww babe, you JUST missed us!! 😭 We were in {C} {D} 💔 We rotate cities "
    "every 1–2 years — but you're locked into our 2027 wishlist for sure 💄✨",
]

# 2026-05-20 PM #2 — Event-interest replies (per Lauren's directive)
# For Bucket B items where customer expressed interest in a specific event
# AND we identified the event from post context. Includes venue + dates + cute closing.
_INTERESTED_LIVE_TEMPLATES = [
    "Yesss come thru!! 💄✨ We're LIVE in {city}, {state} RIGHT NOW 📍 {address}. Open till 5pm today. Can't wait to see you 💕",
    "OMG yay 🥳 We're happening NOW in {city}, {state} 📍 {address} (at {venue}). Til 5pm today — come thru gorgeous 💄💕",
    "Babe we're LIVE 🎉 {city}, {state} 📍 {address}. Open today 10am-5pm — would LOVE to see you 💄✨",
]
_INTERESTED_UPCOMING_SOON_TEMPLATES = [
    "Yay so happy you're interested!! 💄 {city}, {state} is coming up THIS weekend — {dates}, Fri-Sun 10am-5pm 📍 {address} (at {venue}). Can't wait to see you ✨💕",
    "Aww yes!! 🎉 We're hitting {city}, {state} {dates}!! 📍 {address}. Fri-Sun 10am-5pm, free entry + parking. Would love to see you there 💄✨",
    "OMG yay 💄 {city}, {state} is THIS weekend!! {dates} 📍 {address}. Mark your calendar — can't wait to see you 💕✨",
]
_INTERESTED_UPCOMING_TEMPLATES = [
    "Yay so happy you're interested!! 💄 {city}, {state} is on for {dates} 📍 {address} (at {venue}). Fri-Sun 10am-5pm, free entry + parking. Can't wait to see you ✨💕",
    "Aww yesss 🎉 {city}, {state} happens {dates}!! 📍 {address}. We'd love to see you there 💄✨",
    "Babe yes!! 💕 We're coming to {city}, {state} {dates} 📍 {address} (at {venue}). Mark it on the cal — can't wait 💄✨",
]
_INTERESTED_PAST_TEMPLATES = [
    "Aww babe, we were just in {city}, {state} {dates} 😭💔 The sale already wrapped — but we'll do our BEST to come back. Stay close on FB for the announcement 💄✨",
    "Oof you JUST missed us!! 😩 {city}, {state} happened {dates}. We rotate cities every 1-2 years — putting {city} on the wishlist for 2027 💄💕",
]

# 2026-05-20 PM #2 — Messenger DM "Where is location?" with NO post context.
# We can't know which event they're asking about, so soften with the
# next-upcoming event + ask which city they meant. Lauren reviews/edits.
# 2026-07-02 — Lauren's directive: general "where is the event?" questions get
# the FULL tour schedule page (direct URL per IRON RULE #6), so nobody has to
# guess which city they meant. Next-event mention kept as a bonus.
SCHEDULE_PAGE_URL = "https://www.themakeupblowout.com/#events"
_LOCATION_NO_CONTEXT_TEMPLATES = [
    "Hey love 💕 All our upcoming cities & dates are right here 👉 " + SCHEDULE_PAGE_URL + " — our next stop is {next_city}, {next_state} {next_dates} 📍 {next_address}. See you there! 💄✨",
    "Hi babe 💜 You can see the FULL 2026 tour schedule here 👉 " + SCHEDULE_PAGE_URL + " ✨ Next up: {next_city}, {next_state} — {next_dates} 📍 {next_address} 💄",
]


def _seeded_pick(pool: list, seed_str: str) -> str:
    """Pick a template deterministically from the conv/comment ID — same input
    gets the same template every run, but different inputs spread across the pool."""
    if not pool:
        return ""
    h = int(_hashlib.md5(seed_str.encode("utf-8")).hexdigest(), 16) if seed_str else 0
    return pool[h % len(pool)]


def _find_in_schedule(place: str, schedule: dict):
    """
    Try to find this place in the schedule with fuzzy matching:
    - Exact lowercase match
    - Substring match (e.g. 'columbia' matches 'columbia, mo')
    - Word-by-word: 'sanbernardino' matches 'san bernardino, ca'
    Returns (city_key, dates) tuple or None.
    """
    p = _normalize_place(place)
    if not p:
        return None
    sched_keys = list(schedule.keys())
    # Exact
    for k in sched_keys:
        if _normalize_place(k) == p:
            return (k, schedule[k])
    # Substring — 2026-07-02 FIX: whole-word matching only + min length 3.
    # The old raw `p in nk` matched "no" inside "re-no-, nv" and replied about
    # Reno to "no texas?" comments.
    import re as _re_sched
    def _word_in(a, b):
        return bool(_re_sched.search(r"(?<![a-z0-9])" + _re_sched.escape(a) + r"(?![a-z0-9])", b))
    if len(p) >= 3:
        for k in sched_keys:
            nk = _normalize_place(k)
            city_part = nk.split(",")[0].strip()
            if _word_in(p, nk) or (len(city_part) >= 3 and _word_in(city_part, p)):
                return (k, schedule[k])
    # Concatenated (sanbernardino → san bernardino) — min 5 chars (was 2:
    # "no" matched "northlasvegas")
    p_compact = p.replace(" ", "")
    if len(p_compact) >= 5:
        for k in sched_keys:
            if _normalize_place(k).replace(" ", "").startswith(p_compact[:8]):
                return (k, schedule[k])
    return None


def city_rotation_reply(place: str, *, seed: str = "") -> str:
    """
    Off-schedule reply, varied per conversation. Pulls from
    _OFF_SCHEDULE_TEMPLATES, picked deterministically from `seed` (conv_id).
    """
    p = _prettify_place(place) if place else "your area"
    template = _seeded_pick(_OFF_SCHEDULE_TEMPLATES, seed or p)
    return template.format(C=p)


def city_past_reply(city_title: str, dates: str, *, seed: str = "") -> str:
    """Past-event reply — 'we were just there {dates}'."""
    template = _seeded_pick(_PAST_SCHEDULE_TEMPLATES, seed or city_title)
    return template.format(C=city_title, D=dates)


def city_on_schedule_reply(city_title: str, dates: str, *, seed: str = "") -> str:
    """Future/upcoming on-schedule reply, varied per conversation."""
    template = _seeded_pick(_ON_SCHEDULE_TEMPLATES, seed or city_title)
    return template.format(C=city_title, D=dates)


# Tag-only pattern: 2-5 proper-cased names separated by spaces, nothing else
# (matches "Marilyn Rodriguez Sandy Rodriguez" / "John Smith" — friend tags on FB)
import re as _re_classify
TAG_ONLY_PAT = _re_classify.compile(r"^[A-ZA-Za-zÀ-ÿ\u00C0-\u017F]+(\s+[A-ZA-Za-zÀ-ÿ\u00C0-\u017F]+){1,5}\s*$")


def _is_emoji_only_friendly(text: str) -> bool:
    """True if text is ONLY emojis AND none of them are negative/angry.

    Lauren's directive 2026-05-14 PM: "כשיש רק אימוג'י חמודים ולא רעים תעשה
    DONE. אל תעלה אותם זה מיותר". Comments like "🙈🙈🙈" or "💜💜" should be
    auto-marked handled — no need to bother Lauren.

    Negative emoji set (anything that suggests displeasure / anger / complaint)
    will FAIL this check and the comment will still surface to Lauren.
    """
    if not text:
        return False
    t = text.strip()
    if not t:
        return False
    # Strip whitespace + zero-width chars
    import unicodedata
    cleaned = "".join(ch for ch in t if ch.strip())
    if not cleaned:
        return False
    # Negative emojis — if ANY appear, return False so the comment surfaces
    negative_emojis = {
        "😡","🤬","😠","😤","😣","😖","😞","😔","😟","😕","☹️","🙁","😢","😭","😩","😫",
        "💔","🖕","👎","💩","🤢","🤮","😒","😑","😐","🙄","😏","😈","👿","💀","☠️",
    }
    if any(neg in cleaned for neg in negative_emojis):
        return False
    # Check if every character is an emoji / symbol / whitespace
    for ch in cleaned:
        cat = unicodedata.category(ch)
        # Emoji ranges + symbol categories
        codepoint = ord(ch)
        is_emoji = (
            cat.startswith("S")  # Symbol categories: So, Sm, Sk, Sc
            or (0x1F300 <= codepoint <= 0x1FAFF)
            or (0x2600 <= codepoint <= 0x27BF)
            or (0xFE00 <= codepoint <= 0xFE0F)  # variation selectors
            or (0x1F1E6 <= codepoint <= 0x1F1FF)  # regional indicators (flags)
            or codepoint == 0x200D  # zero-width joiner (combines emojis)
            or codepoint == 0x200C
        )
        if not is_emoji and not ch.isspace():
            return False
    return True


def _is_reaction_or_empty(text: str) -> bool:
    """Detect the placeholder text the script writes when Meta API returns
    a conversation with no real message body (just a reaction emoji or
    attachment-only). Per SKILL.md 'Auto-move to Done': skip these."""
    if not text:
        return True
    t = text.lower().strip()
    markers = [
        "(reaction or empty message)",
        "(no text — perhaps an attachment)",
        "(empty message)",
        "(deleted)",
    ]
    return any(m in t for m in markers)


def _is_positive_closer(text: str) -> bool:
    """Detect short positive messages with no question — 'thanks!', 'love you all',
    'great time last year'. Per SKILL.md 'Auto-move to Done'.

    Heuristics:
    - Length <= 80 chars
    - No '?' (no question)
    - Has at least one positive keyword
    - No city/place mention pattern (those need real reply)
    - Not asking about scheduling
    """
    if not text:
        return False
    t = text.strip()
    if len(t) > 80:
        return False
    if "?" in t:
        return False
    t_lower = t.lower()
    positives = [
        "thank you","thanks","thx","ty","love","loved","amazing","great","perfect",
        "awesome","wonderful","fire","best","grateful","mean a lot","appreciate",
        "תודה","אהבתי","אוהבת","מעולה","מדהים","תענוג",
    ]
    has_pos = any(p in t_lower for p in positives)
    if not has_pos:
        return False
    # Reject only if it contains CLEAR scheduling-question intent.
    # "had a great time last year" is positive-closer, NOT scheduling.
    # Trim list to unambiguous scheduling signals only.
    schedule_words = [
        "when ", " when","where","coming back","come back","please come",
        "going to","be in","near me","around here","next event","next sale",
        "מתי","איפה",
    ]
    if any(w in t_lower for w in schedule_words):
        return False
    return True


# US state abbreviation → full name (for short messages like "Ca" / "TX")
_US_STATES = {
    "AL":"Alabama","AK":"Alaska","AZ":"Arizona","AR":"Arkansas","CA":"California",
    "CO":"Colorado","CT":"Connecticut","DE":"Delaware","FL":"Florida","GA":"Georgia",
    "HI":"Hawaii","ID":"Idaho","IL":"Illinois","IN":"Indiana","IA":"Iowa",
    "KS":"Kansas","KY":"Kentucky","LA":"Louisiana","ME":"Maine","MD":"Maryland",
    "MA":"Massachusetts","MI":"Michigan","MN":"Minnesota","MS":"Mississippi","MO":"Missouri",
    "MT":"Montana","NE":"Nebraska","NV":"Nevada","NH":"New Hampshire","NJ":"New Jersey",
    "NM":"New Mexico","NY":"New York","NC":"North Carolina","ND":"North Dakota","OH":"Ohio",
    "OK":"Oklahoma","OR":"Oregon","PA":"Pennsylvania","RI":"Rhode Island","SC":"South Carolina",
    "SD":"South Dakota","TN":"Tennessee","TX":"Texas","UT":"Utah","VT":"Vermont",
    "VA":"Virginia","WA":"Washington","WV":"West Virginia","WI":"Wisconsin","WY":"Wyoming",
}


# Reverse lookup: lowercase state name -> code (e.g. "colorado" -> "CO").
_STATE_NAME_TO_CODE = {name.lower(): code for code, name in _US_STATES.items()}

# Words that are NOT part of a city name when they precede a state token.
_PLACE_STOPWORDS = {
    "in","to","at","near","around","for","a","an","one","do","does","did","done",
    "the","is","are","you","u","come","coming","came","when","we","can","could",
    "would","will","consider","considering","have","having","host","hosting","bring",
    "bringing","throw","plan","planning","plans","hoping","hope","wish","want","wanna",
    "love","please","pls","plz","hey","hi","of","your","our","next","back","go","going",
    "gonna","visit","visiting","make","made","or","and","y'all","yall","sale","event",
    "show","popup","pop-up","any","chance","time","ever","also","too","this","that",
    # 2026-07-02 — question/filler words that were leaking into place names
    # ("What about Texas" → city "What about"; "no texas" → city "no";
    #  "add San Antonio" → city "add San Antonio")
    "no","not","yes","what","about","add","adding","why","where","how","who",
    "never","again","soon","need","needs","here","there","us","only","even",
    "still","yet","miss","missing","skip","skipped",
    # 2026-07-02 #2 — "Im Washington" → city was captured as "year Im"
    "im","i'm","year","years","month","months","week","weekend","today",
    "sad","happy","glad","live","living","stay","staying","from","born",
    "see","saw","dont","don't","didnt","didn't","wont","won't",
}

# 2026-06-01 — Event-request phrasing WITHOUT the usual "when/coming/visit" trigger
# words. Catches "can y'all consider do one in <City>", "any chance you bring a sale
# to <City>", "you should host an event in <City>". Requires an event-object word
# (one/sale/event/show/pop-up/stop) so it won't false-fire on "do you have it in stock".
EVENT_REQUEST_PAT = re.compile(
    r"\b(?:consider|considering|do|host|hosting|have|having|bring|bringing|throw|plan|planning|"
    r"hoping|hope|wish|want|wanna|y'?all|chance|could\s+you|should)\b"
    r"[^.?!]*?\b(?:one|a\s+sale|an?\s+event|sale|event|show|pop-?up|stop)\b"
    r"[^.?!]*?\b(?:in|to|at|near|around|for)\s+"
    r"([A-Za-z][A-Za-z0-9.\-']*(?:\s+[A-Za-z][A-Za-z0-9.\-']*){0,3})",
    re.IGNORECASE,
)

# 2026-06-01 — "I'm waiting for / hoping for <place>" interest phrasing. Gated:
# the captured place must validate (state token or schedule hit) before we auto-
# reply, so "waiting for my order" stays with Lauren. Colloquial regions like
# "high desert" are handled separately by _detect_known_region.
WAITING_PLACE_PAT = re.compile(
    r"\b(?:wait(?:ing)?\s+(?:for|on)|hoping\s+for|praying\s+for|dying\s+for|"
    r"holding\s+out\s+for|can'?t\s+wait\s+for)\b\s+(?:the\s+|you\s+to\s+come\s+to\s+|an?\s+event\s+in\s+)?"
    r"([A-Za-z][A-Za-z0-9.\-']*(?:\s+[A-Za-z][A-Za-z0-9.\-']*){0,3})",
    re.IGNORECASE,
)

# Colloquial US region / metro nicknames people ask about (normalized, no leading "the").
_KNOWN_REGIONS = {
    "high desert","low desert","the desert","victor valley","antelope valley",
    "coachella valley","imperial valley","inland empire","san fernando valley",
    "san gabriel valley","san joaquin valley","central valley","bay area","east bay",
    "south bay","north bay","central coast","gold coast","gulf coast","treasure valley",
    "magic valley","hill country","low country","tri cities","tri state","quad cities",
    "pacific northwest","inland northwest","southern california","northern california",
    "central california","socal","norcal","the valley","high country","palm springs area",
}


def _detect_known_region(text: str):
    """Return a prettified region name (e.g. 'High Desert') if a known colloquial
    region nickname appears in the text, else None."""
    if not text:
        return None
    norm = re.sub(r"[^a-z0-9 ]+", " ", text.lower())
    norm = " " + re.sub(r"\s+", " ", norm).strip() + " "
    for region in sorted(_KNOWN_REGIONS, key=len, reverse=True):
        if (" " + region + " ") in norm:
            return region.title()
    return None


def _place_with_state(text: str):
    """Detect a 'City State' / 'City, ST' location mention anywhere in the text.

    Handles organic-comment shapes that name a place + its state but lack a
    question verb the older patterns required, e.g.:
        '? Denver Colorado'                       -> ('Denver', 'CO')
        'Abilene TX?'                             -> ('Abilene', 'TX')
        "Can y'all consider do one in Abilene TX" -> ('Abilene', 'TX')
        'South Denver, CO'                        -> ('South Denver', 'CO')

    Tokenize, find a state token (full name case-insensitive; or a 2-letter code
    that is UPPERCASE to avoid colliding with words like 'in'/'or'/'oh'), then
    walk backwards collecting up to 3 city words until a stopword. Returns
    (city, state_code) or None. A bare state with no city returns None.
    """
    if not text:
        return None
    raw = re.findall(r"[A-Za-z][A-Za-z.'\-]*", text)
    if not raw:
        return None
    low = [w.lower() for w in raw]
    n = len(raw)
    found = None  # (code, start_idx)
    i = 0
    while i < n:
        if i + 1 < n:
            two = low[i] + " " + low[i + 1]
            if two in _STATE_NAME_TO_CODE:
                found = (_STATE_NAME_TO_CODE[two], i)
                i += 2
                continue
        if low[i] in _STATE_NAME_TO_CODE:
            found = (_STATE_NAME_TO_CODE[low[i]], i)
        elif len(raw[i]) == 2 and raw[i].isupper() and raw[i] in _US_STATES:
            found = (raw[i], i)
        i += 1
    if not found:
        return None
    code, sidx = found
    city_words = []
    j = sidx - 1
    while j >= 0 and len(city_words) < 3:
        if low[j] in _PLACE_STOPWORDS:
            break
        if low[j] in _STATE_NAME_TO_CODE or (len(raw[j]) == 2 and raw[j].isupper() and raw[j] in _US_STATES):
            break
        city_words.insert(0, raw[j])
        j -= 1
    if not city_words:
        return None
    return (" ".join(city_words), code)



_MONTH_NAMES = {"january","february","march","april","may","june","july","august",
           "september","october","november","december",
           "jan","feb","mar","apr","jun","jul","aug","sep","sept","oct","nov","dec"}
# Trailing words to peel off a captured place ("Denver please" -> "Denver").
_PLACE_TRAIL_TRIM = {"please","pls","plz","soon","again","next","this","that","too",
                     "thanks","thank","thx","ty","tho","though","ever","time","maybe",
                     "yet","now","here","there","ok","okay","already","still","guys","girls"}
# Words that are clearly NOT a city even though grammar lets them slip through.
_NON_PLACE_WORDS = {"october","stock","town","here","there","today","tomorrow",
                    "soon","future","general","store","online","person","time","start",
                    "yet","stuff","things","everything","anything","makeup","lipstick",
                    "products","product","refund","order","shipment","package","delivery",
                    "response","reply","restock","email","website","site","sale","event"}


def _clean_place_candidate(cand: str) -> str:
    """Trim trailing filler words off a captured place phrase. Returns '' if nothing left."""
    if not cand:
        return ""
    words = cand.strip().rstrip("?.!,").split()
    while words and words[-1].lower() in _PLACE_TRAIL_TRIM:
        words.pop()
    # 2026-07-02 — strip leading filler/verbs ("add San Antonio", "please come to X")
    _LEAD_TRIM = _PLACE_STOPWORDS | {"see", "saw", "dont", "don't", "didnt", "didn't"}
    while words and words[0].lower() in _LEAD_TRIM:
        words.pop(0)
    # 2026-07-02 #2 — strip a trailing state token ("Austin Tx" → "Austin",
    # "Seattle wa" → "Seattle"); keep single-word candidates intact ("Texas").
    if len(words) > 1:
        lastw = words[-1].rstrip("?.!,")
        if (len(lastw) == 2 and lastw.upper() in _US_STATES) or lastw.lower() in _STATE_NAME_TO_CODE:
            words.pop()
    return " ".join(words).strip()


def _is_non_place(place: str) -> bool:
    """True if `place` is obviously not a city (a month name, filler word, etc.).
    Used to suppress off-schedule rotation replies for false-positive captures
    like 'October' in 'I'm coming in October'."""
    if not place:
        return True
    p = place.strip().lower().rstrip("?.!,")
    if p in _NON_PLACE_WORDS:
        return True
    # 2026-07-02 — every word is a stopword ("what about", "no", "add") = not a place
    if all(w in _PLACE_STOPWORDS for w in p.split()):
        return True
    # 2026-07-02 #2 — any component word is a known non-place ("makeup show with beans")
    if any(w in _NON_PLACE_WORDS for w in p.split()):
        return True
    last = p.split()[-1] if p.split() else ""
    if last in _MONTH_NAMES or p in _MONTH_NAMES:
        return True
    return False


def _detect_state_name(text: str) -> str:
    """Return state code (e.g. 'CA') if the text is essentially just a state
    reference. Returns None if not a clean state mention.

    Examples that match: 'Ca', 'CA', 'California', 'california?', 'Ca?'.
    Examples that don't: 'CA event when?', 'Going to CA next month'.
    """
    if not text:
        return None
    t = text.strip().rstrip("?").strip()
    if len(t) > 15:
        return None
    # Match by abbreviation (case-insensitive)
    upper = t.upper()
    if upper in _US_STATES:
        return upper
    # Match by full state name
    for code, name in _US_STATES.items():
        if t.lower() == name.lower():
            return code
    return None




def _is_tag_only(text: str) -> bool:
    """True if text appears to be JUST friend-tag names with no real content.

    Heuristics:
    - Total length < 60 chars (longer = probably real text)
    - All words are 1+ capital letter starts (proper-name style)
    - No punctuation except periods/apostrophes
    - No question words, lowercase verbs, or sentence-like patterns
    """
    if not text:
        return False
    t = text.strip()
    if len(t) > 60:
        return False
    # Reject if there's a question mark, exclamation, or common lowercase words
    if any(p in t for p in ["?", "!", "$", "@"]):
        return False
    lower_only_words = ["the","is","are","you","i","when","where","how","what","why","do","can","will","please","thanks","thank","need","want","got"]
    words_lower = t.lower().split()
    if any(w in lower_only_words for w in words_lower):
        return False
    # Each word should start with uppercase (allows accented characters)
    words = t.split()
    if len(words) < 2 or len(words) > 6:
        return False
    if not all(w[0].isupper() for w in words if w):
        return False
    return True



# 2026-07-02 — State-aware schedule lookup. "Are you coming to New Mexico??"
# was getting a rotation ("not this year") reply even though Santa Fe + Albuquerque
# ARE on the 2026 schedule. When a mention resolves to a STATE, answer with that
# state's upcoming events; when a city misses but its state has events, say so.
_STATE_ON_SCHEDULE_TEMPLATES = [
    "YESSS 🎉 We ARE coming to {S}: {LIST}! Fri–Sun 10am–5pm, free admission 💄 See you there?",
    "Girl, good news — {S} is ON the 2026 tour: {LIST} 💕 Free entry, 10am–5pm. Mark your calendar!",
    "We got you, {S}! 🙌 Catch us at {LIST} — free admission, Fri–Sun 10am–5pm 💄",
]
_CITY_MISS_STATE_HIT_TEMPLATES = [
    "{C} isn't on the 2026 stop list 💔 BUT we ARE coming to {S}: {LIST}! Worth the trip — free admission, 10am–5pm 💄",
    "No {C} this year 🥺 but don't worry — {S} is on the 2026 tour: {LIST}! Free entry, Fri–Sun 10am–5pm 💕",
]


def _as_state_code(place: str):
    """'New Mexico' / 'NM' / 'nevada' → 'NM'/'NV'; None if not a state."""
    if not place:
        return None
    t = place.strip().rstrip("?.!,").strip()
    if t.upper() in _US_STATES and len(t) == 2:
        return t.upper()
    return _STATE_NAME_TO_CODE.get(t.lower())


def _state_schedule_hits(state_code: str, schedule: dict, max_hits: int = 2,
                         past: bool = False) -> list:
    """Schedule entries in this state → [(city_title, dates)]. past=False →
    upcoming only; past=True → already-happened only."""
    hits = []
    sc = (state_code or "").lower()
    for k, d in schedule.items():
        if "," in k and k.split(",")[1].strip().lower() == sc and _is_past(d) == past:
            hits.append((k.split(",")[0].strip().title(), d))
            if len(hits) >= max_hits:
                break
    return hits


_STATE_PAST_TEMPLATES = [
    "Aww babe, we already hit {S} this year — {LIST} 💔 That wraps {S} for 2026, but it's going straight on the 2027 wishlist 📝💄 Follow our FB for the announcement 👀",
    "We were JUST in {S}: {LIST} 😭 No more {S} stops this year — but we rotate back every 1–2 years, so stay close 💕",
]


def _fmt_state_hits(hits: list) -> str:
    return " & ".join(f"{c} ({d})" for c, d in hits)


def _classify_state_mention(state_code: str, kb: dict, seed: str) -> dict:
    """Bucket-A reply for a state-level mention."""
    state_name = _US_STATES.get(state_code, state_code)
    hits = _state_schedule_hits(state_code, kb.get("schedule", {}))
    if hits:
        tmpl = _seeded_pick(_STATE_ON_SCHEDULE_TEMPLATES, seed + state_code)
        return {"bucket": "A",
                "reason": f"State mention ({state_name}) — {len(hits)} upcoming event(s) in state",
                "reply": tmpl.format(S=state_name, LIST=_fmt_state_hits(hits)),
                "event_chip": {"city": hits[0][0], "state": state_code,
                               "dates": hits[0][1], "status": "upcoming", "confidence": "high"}}
    past_hits = _state_schedule_hits(state_code, kb.get("schedule", {}), max_hits=3, past=True)
    if past_hits:
        tmpl = _seeded_pick(_STATE_PAST_TEMPLATES, seed + state_code)
        return {"bucket": "A",
                "reason": f"State mention ({state_name}) — only past 2026 events in state",
                "reply": tmpl.format(S=state_name, LIST=_fmt_state_hits(past_hits))}
    return {"bucket": "A",
            "reason": f"State mention ({state_name}) — no 2026 events in state, rotation reply",
            "reply": city_rotation_reply(state_name, seed=seed + state_code)}


# 2026-07-02 — "what about <place>?" phrasing ("What about San Diego?" was Bucket B)
WHAT_ABOUT_PAT = re.compile(
    r"\bwhat\s+about\s+([A-Za-z][A-Za-z.\-']*(?:\s+[A-Za-z][A-Za-z.\-']*){0,3})",
    re.IGNORECASE,
)

def classify(text: str, kb: dict) -> dict:
    """
    Return {bucket: 'A'|'B'|'NEG'|'SKIP', reason, suggested_reply} for a message.
    Bucket A = answerable from KB (auto-reply).
    Bucket B = needs Lauren (low confidence, complex, or no KB match).
    NEG     = never auto-reply (negative/skeptical comment).
    SKIP    = no action needed (tag-only friend mentions, etc.) — never auto-reply, never shown to Lauren.
    """
    if not text or len(text.strip()) < 2:
        return {"bucket": "SKIP", "reason": "empty/too short (auto-done)", "reply": None}

    # 2026-06-01 — detect a "City State" / "City, ST" location mention up front so
    # it can rescue messages that older checks would mis-bucket (a bare "Denver
    # Colorado" looks like a friend-tag; "Abilene TX" has no question verb). The
    # actual reply is built lower down, after the NEGATIVE / FAQ checks.
    place_state = _place_with_state(text)

    # Tag-only friend mentions (e.g. "Marilyn Rodriguez Sandy Rodriguez")
    if _is_tag_only(text) and not place_state:
        return {"bucket": "SKIP", "reason": "tag-only-friends (no question)", "reply": None}

    # 2026-05-14 PM — emoji-only friendly comments → auto-handled, never shown to Lauren
    if _is_emoji_only_friendly(text):
        return {"bucket": "SKIP", "reason": "emoji-only-friendly (auto-done)", "reply": None}

    # 2026-05-14 PM — reaction or empty-body messages from Meta API → auto-done
    if _is_reaction_or_empty(text):
        return {"bucket": "SKIP", "reason": "reaction-or-empty (auto-done)", "reply": None}

    # 2026-05-14 PM — positive closer with no question → auto-done (per SKILL.md)
    if _is_positive_closer(text) and not place_state:
        return {"bucket": "SKIP", "reason": "positive-closer-no-question (auto-done)", "reply": None}

    # 2026-05-14 PM — state-name-only message ("Ca", "California") → CA event rotation reply
    state_code = _detect_state_name(text)
    if state_code:
        # 2026-07-02 — consult the schedule (was blind rotation reply even when
        # the state HAS upcoming events)
        return _classify_state_mention(state_code, kb, kb.get("_seed", ""))

    t = text.strip()

    # "Where?" comments — only meaningful when we have the post context
    post_ctx = kb.get("_post_context") if isinstance(kb, dict) else None
    if post_ctx and WHERE_PAT.search(t):
        info = summarize_event_from_caption(post_ctx.get("caption", ""))
        if info:
            info = _enrich_event_info(info, kb.get("_venues", []))
            seed = (kb.get("_seed", "") or "") + "where"
            status = event_status(info.get("dates", ""))
            address = info.get("address") or "(address tba)"
            venue   = info.get("venue") or "the venue"
            city, state = info.get("city",""), info.get("state","")
            dates = info.get("dates","")
            if status == "live":
                tmpl = _seeded_pick(_WHERE_LIVE_TEMPLATES, seed)
            elif status == "upcoming-soon":
                tmpl = _seeded_pick(_WHERE_UPCOMING_SOON_TEMPLATES, seed)
            elif status == "upcoming":
                tmpl = _seeded_pick(_WHERE_UPCOMING_TEMPLATES, seed)
            elif status == "past":
                tmpl = _seeded_pick(_WHERE_PAST_TEMPLATES, seed)
            else:
                # 2026-05-20 PM #2 — unknown date format; default to upcoming
                tmpl = _seeded_pick(_WHERE_UPCOMING_TEMPLATES, seed)
                status = "upcoming"
            if tmpl:
                return {"bucket": "A",
                        "reason": f"'Where?' on event reel — status={status}, city={city}",
                        "reply": tmpl.format(city=city, state=state, dates=dates,
                                             address=address, venue=venue),
                        "event_chip": {"city": city, "state": state, "dates": dates,
                                       "status": status, "confidence": "high"}}

    # Negative — never auto-reply
    if NEGATIVE_KEYWORDS.search(t):
        return {"bucket": "NEG", "reason": "negative keywords detected", "reply": None}

    # Past-purchase or customer-issue — always Bucket B (Lauren handles personally)
    if CUSTOMER_ISSUE_PAT.search(t):
        return {"bucket": "B",
                "reason": "Customer issue / past purchase pattern — needs Lauren personally",
                "reply": None}

    # Online-order question — direct her to the website
    if ONLINE_ORDER_PAT.search(t):
        return {"bucket": "A", "reason": "Online order/availability question",
                "reply": _seeded_pick(_ONLINE_ORDER_TEMPLATES, kb.get("_seed","")+"online")}

    # Open today / are you here today
    if OPEN_TODAY_PAT.search(t):
        venues = kb.get("_venues", [])
        live = find_event_today(venues)
        seed = kb.get("_seed", "") + "today"
        if live:
            tmpl = _seeded_pick(_OPEN_TODAY_LIVE_TEMPLATES, seed)
            cf = live["city"]
            city = cf.split(",")[0].strip(); state = cf.split(",")[1].strip() if "," in cf else ""
            return {"bucket": "A", "reason": f"Open today? — LIVE at {cf}",
                    "reply": tmpl.format(city=city, state=state, address=live["address"], venue=live["venue"], dates=live["dates"])}
        nxt = find_next_event(venues)
        if nxt:
            tmpl = _seeded_pick(_OPEN_TODAY_CLOSED_TEMPLATES, seed)
            cf = nxt["city"]
            city = cf.split(",")[0].strip(); state = cf.split(",")[1].strip() if "," in cf else ""
            return {"bucket": "A", "reason": f"Open today? — closed; next is {cf}",
                    "reply": tmpl.format(city=city, state=state, address=nxt["address"], venue=nxt["venue"], dates=nxt["dates"])}

    if PAYMENT_PAT.search(t):
        return {"bucket": "A", "reason": "Payment methods",
                "reply": _seeded_pick(_PAYMENT_TEMPLATES, kb.get("_seed","")+"pay")}
    if PARKING_PAT.search(t):
        return {"bucket": "A", "reason": "Parking",
                "reply": _seeded_pick(_PARKING_TEMPLATES, kb.get("_seed","")+"park")}
    if HOURS_PAT.search(t):
        return {"bucket": "A", "reason": "Hours",
                "reply": _seeded_pick(_HOURS_TEMPLATES, kb.get("_seed","")+"hrs")}
    if KIDS_PAT.search(t):
        return {"bucket": "A", "reason": "Kids welcome",
                "reply": _seeded_pick(_KIDS_TEMPLATES, kb.get("_seed","")+"kid")}

    # Address question (only with post context)
    post_ctx2 = kb.get("_post_context") if isinstance(kb, dict) else None
    if post_ctx2 and ADDRESS_PAT.search(t):
        info = summarize_event_from_caption(post_ctx2.get("caption", ""))
        if info:
            info = _enrich_event_info(info, kb.get("_venues", []))
        if info and info.get("address"):
            seed = (kb.get("_seed","") or "") + "addr"
            status = event_status(info.get("dates",""))
            tmpl = None
            if status == "live":          tmpl = _seeded_pick(_WHERE_LIVE_TEMPLATES, seed)
            elif status == "upcoming-soon": tmpl = _seeded_pick(_WHERE_UPCOMING_SOON_TEMPLATES, seed)
            elif status == "upcoming":      tmpl = _seeded_pick(_WHERE_UPCOMING_TEMPLATES, seed)
            elif status == "past":          tmpl = _seeded_pick(_WHERE_PAST_TEMPLATES, seed)
            if tmpl:
                return {"bucket": "A", "reason": f"Address — status={status}",
                        "reply": tmpl.format(city=info.get("city",""), state=info.get("state",""),
                                             dates=info.get("dates",""), address=info["address"],
                                             venue=info.get("venue") or "the venue")}

    # 2026-07-02 #2 — influencer / collab / PR inquiries must reach Lauren
    # (a "local content creator" message was getting the brands-list FAQ reply)
    if re.search(r"\b(content\s+creator|influencer|collab(?:oration)?|partnership|"
                 r"pr\s+list|brand\s+ambassador|work\s+(?:with|together)|promote\s+you)\b", t, re.IGNORECASE):
        return {"bucket": "B", "reason": "Influencer/collab inquiry — Lauren personally", "reply": None}

    # FAQ keyword match (priority over city question — "are you free?" outranks "are you coming")
    for pattern, faq_q in FAQ_KEYWORDS:
        if pattern.search(t):
            if faq_q:
                for f in kb["faqs"]:
                    if f["q"].lower().startswith(faq_q.lower()[:20]):
                        return {"bucket": "A", "reason": f"FAQ match: {f['q']}",
                                "reply": f["a"]}
            return {"bucket": "B", "reason": "keyword matched but no KB answer",
                    "reply": None}

    # 2026-05-20 PM #2 — "Interested in this event" / "i'm going" with post context.
    # Generates event-specific draft (Bucket B — Lauren reviews) with venue + dates.
    post_ctx_int = kb.get("_post_context") if isinstance(kb, dict) else None
    if post_ctx_int and INTERESTED_PAT.search(t):
        info = summarize_event_from_caption(post_ctx_int.get("caption", ""))
        if info:
            info = _enrich_event_info(info, kb.get("_venues", []))
            status = event_status(info.get("dates",""))
            seed = (kb.get("_seed","") or "") + "interested"
            if status == "live":          tmpl = _seeded_pick(_INTERESTED_LIVE_TEMPLATES, seed)
            elif status == "upcoming-soon": tmpl = _seeded_pick(_INTERESTED_UPCOMING_SOON_TEMPLATES, seed)
            elif status == "upcoming":      tmpl = _seeded_pick(_INTERESTED_UPCOMING_TEMPLATES, seed)
            elif status == "past":          tmpl = _seeded_pick(_INTERESTED_PAST_TEMPLATES, seed)
            else:
                # status="unknown" — caption's date format was unparseable, but
                # city/state/venue are valid. 2026-07-02 #2: if the POST itself is
                # old (>21 days), the event almost certainly passed — replying
                # "we're coming!" months later reads bot-broken. Skip instead.
                _post_date = (kb.get("_post_context") or {}).get("date", "")
                try:
                    _age_days = (dt.date.today() - dt.datetime.fromisoformat(
                        _post_date.replace("Z", "+00:00")).date()).days
                except Exception:
                    _age_days = 0
                if _age_days > 21:
                    return {"bucket": "SKIP",
                            "reason": "interest on old post, event date unparseable — no reply needed",
                            "reply": None}
                tmpl = _seeded_pick(_INTERESTED_UPCOMING_TEMPLATES, seed)
                status = "upcoming"
            if tmpl:
                reply = tmpl.format(
                    city=info.get("city",""), state=info.get("state",""),
                    dates=info.get("dates",""), address=info.get("address") or "(address TBA)",
                    venue=info.get("venue") or "the venue",
                )
                # 2026-05-20 PM #3 — Lauren's directive: "תידע לענות איך שהיא מקבלת את ההודעה
                # ושלא תשלח לי לאישור אלא כמה שפחות". Flipped B→A: high-confidence event
                # match with post_ctx is safe to auto-reply (post identifies the event).
                # CRITICAL: auto-reply also gets us inside Meta's 24-hour Messenger window
                # — items that sit in Bucket B for 4+ days can no longer be replied to.
                return {"bucket": "A",
                        "reason": f"Event interest auto-reply — {info['city']}, {info['state']} ({status})",
                        "reply": reply,
                        "event_chip": {"city": info.get("city",""), "state": info.get("state",""),
                                       "dates": info.get("dates",""), "status": status,
                                       "confidence": "high"}}

    # 2026-06-01 — City+State mention (or event-request phrasing) detected up front.
    # This is the fix for organic comments like "? Denver Colorado" and
    # "Can y'all consider do one in Abilene TX" that name a place but lack the
    # verbs the older CITY_QUESTION_PAT required. Look the city up in the schedule
    # and answer on/past/off-schedule per the KB City Rotation Policy.
    _place_for_reply = None
    if place_state:
        _place_for_reply = place_state[0]
    else:
        m_evt = EVENT_REQUEST_PAT.search(t) or WHAT_ABOUT_PAT.search(t)
        if m_evt:
            cand = _clean_place_candidate(m_evt.group(1))
            # strip a trailing state token off the captured place ("Abilene TX" -> "Abilene")
            ps2 = _place_with_state(cand)
            _place_for_reply = ps2[0] if ps2 else cand
        # Colloquial region nickname ("high desert", "inland empire", ...)
        if not _place_for_reply:
            reg = _detect_known_region(t)
            if reg:
                _place_for_reply = reg
        # "waiting for / hoping for <place>" — gated: only auto-reply if the place
        # validates (state token or schedule hit); otherwise leave to Lauren.
        if not _place_for_reply:
            m_wait = WAITING_PLACE_PAT.search(t)
            if m_wait:
                cand = _clean_place_candidate(m_wait.group(1))
                ps3 = _place_with_state(cand)
                if ps3:
                    _place_for_reply = ps3[0]
                elif cand and _find_in_schedule(cand, kb["schedule"]):
                    _place_for_reply = cand
    if _place_for_reply and _is_non_place(_place_for_reply):
        _place_for_reply = None
    if _place_for_reply:
        seed = kb.get("_seed", "") + _place_for_reply
        # 2026-07-02 — the "place" may actually be a STATE ("What about Texas")
        _st = _as_state_code(_place_for_reply)
        if _st:
            return _classify_state_mention(_st, kb, kb.get("_seed", ""))
        hit = _find_in_schedule(_place_for_reply, kb["schedule"])
        if hit:
            city_key, dates = hit
            city_title = city_key.split(",")[0].strip().title()
            _state = city_key.split(",")[1].strip().upper() if "," in city_key else ""
            if _is_past(dates):
                return {"bucket": "A",
                        "reason": f"City+state mention — '{_place_for_reply}' was on schedule, past ({dates})",
                        "reply": city_past_reply(city_title, dates, seed=seed),
                        "event_chip": {"city": city_title, "state": _state,
                                       "dates": dates, "status": "past", "confidence": "high"}}
            return {"bucket": "A",
                    "reason": f"City+state mention — '{_place_for_reply}' matches upcoming schedule '{city_key}'",
                    "reply": city_on_schedule_reply(city_title, dates, seed=seed),
                    "event_chip": {"city": city_title, "state": _state,
                                   "dates": dates, "status": "upcoming", "confidence": "high"}}
        # 2026-07-02 — city not on schedule, but if we KNOW the state and the
        # state HAS upcoming events, say so ("Las Cruces? no — but Santa Fe &
        # Albuquerque, NM are coming!") instead of a plain rotation reply.
        _st2 = place_state[1] if place_state else None
        if _st2:
            _hits2 = _state_schedule_hits(_st2, kb.get("schedule", {}))
            if _hits2:
                _tmpl2 = _seeded_pick(_CITY_MISS_STATE_HIT_TEMPLATES, seed)
                return {"bucket": "A",
                        "reason": f"City+state mention — '{_place_for_reply}' not on schedule but {_st2} has {len(_hits2)} event(s)",
                        "reply": _tmpl2.format(C=_prettify_place(_place_for_reply),
                                               S=_US_STATES.get(_st2, _st2),
                                               LIST=_fmt_state_hits(_hits2)),
                        "event_chip": {"city": _hits2[0][0], "state": _st2,
                                       "dates": _hits2[0][1], "status": "upcoming",
                                       "confidence": "high"}}
        return {"bucket": "A",
                "reason": f"City+state mention — '{_place_for_reply}' NOT on 2026 schedule, rotation-policy reply",
                "reply": city_rotation_reply(_place_for_reply, seed=seed)}

    # City question intent — "when are you coming to X" / "sale in X"
    m_q = CITY_QUESTION_PAT.search(t)
    # 2026-05-20 — don't false-match BARE_PLACE_PAT on question phrases
    # (was matching "Where is location?" as a place name).
    _starts_with_qword = bool(re.match(r"^\s*(where|when|how|what|why|who|which|can\s+you|do\s+you)\b", t, re.IGNORECASE))
    m_b = BARE_PLACE_PAT.match(t) if (not m_q and not _starts_with_qword) else None
    if m_q or m_b:
        place = _clean_place_candidate(m_q.group(2) if m_q else m_b.group(1))
        # 2026-07-02 — "are you coming to New Mexico??" — the place is a STATE
        _stq = _as_state_code(place)
        if _stq:
            return _classify_state_mention(_stq, kb, kb.get("_seed", ""))
        hit = _find_in_schedule(place, kb["schedule"])
        seed = kb.get("_seed", "") + place  # conv-id + place for variation seed
        if hit:
            city_key, dates = hit
            city_title = city_key.split(",")[0].strip().title()
            if _is_past(dates):
                return {"bucket": "A",
                        "reason": f"City question — '{place}' was on schedule but already happened ({dates})",
                        "reply": city_past_reply(city_title, dates, seed=seed)}
            return {"bucket": "A",
                    "reason": f"City question — '{place}' matches upcoming schedule entry '{city_key}'",
                    "reply": city_on_schedule_reply(city_title, dates, seed=seed)}
        # Strong city-question intent but not on schedule → off-schedule reply
        if m_q and not _is_non_place(place):
            return {"bucket": "A",
                    "reason": f"City question — '{place}' NOT on 2026 schedule, off-schedule reply",
                    "reply": city_rotation_reply(place, seed=seed)}
        # Bare-place pattern (no verb) and not on schedule = ambiguous — Lauren reviews
        return {"bucket": "B",
                "reason": f"Bare-place pattern matched '{place}' but not on schedule — ambiguous, Lauren reviews",
                "reply": None}

    # 2026-05-14 PM — relaxed place mention (informal "come back to X") → city reply
    m_relax = RELAXED_PLACE_PAT.search(t)
    if m_relax:
        place = m_relax.group(1).strip()
        if 1 < len(place) <= 40:
            hit = _find_in_schedule(place, kb["schedule"])
            seed = kb.get("_seed","") + place
            if hit:
                city_key, dates = hit
                city_title = city_key.split(",")[0].strip().title()
                # 2026-05-20 — derive state from city_key for the chip
                _state = city_key.split(",")[1].strip().upper() if "," in city_key else ""
                if _is_past(dates):
                    return {"bucket": "A",
                            "reason": f"Place mention ({place}) — was on schedule, past",
                            "reply": city_past_reply(city_title, dates, seed=seed),
                            "event_chip": {"city": city_title, "state": _state,
                                           "dates": dates, "status": "past", "confidence": "high"}}
                return {"bucket": "A",
                        "reason": f"Place mention ({place}) — on upcoming schedule",
                        "reply": city_on_schedule_reply(city_title, dates, seed=seed),
                        "event_chip": {"city": city_title, "state": _state,
                                       "dates": dates, "status": "upcoming", "confidence": "high"}}
            return {"bucket": "A",
                    "reason": f"Place mention ({place}) — off-schedule rotation reply",
                    "reply": city_rotation_reply(place, seed=seed)}

    # 2026-05-20 PM #2 — Messenger WHERE_PAT without post context.
    # We can't pinpoint the event, so soften with next-upcoming + ask which city.
    # Bucket B (Lauren reviews) with event_chip showing low-confidence next-event guess.
    if WHERE_PAT.search(t) or re.search(r"\b(where|location|address|directions?|map)\b", t, re.IGNORECASE):
        venues = kb.get("_venues", [])
        nxt = find_next_event(venues)
        if nxt:
            cf = nxt["city"]
            nxt_city = cf.split(",")[0].strip()
            nxt_state = cf.split(",")[1].strip() if "," in cf else ""
            seed = (kb.get("_seed","") or "") + "where-no-ctx"
            tmpl = _seeded_pick(_LOCATION_NO_CONTEXT_TEMPLATES, seed)
            reply = tmpl.format(
                first_word="there",
                next_city=nxt_city, next_state=nxt_state,
                next_dates=nxt.get("dates",""), next_address=nxt.get("address","")
            )
            return {"bucket": "A",
                    "reason": f"Location question w/o context — schedule page + next event ({nxt_city})",
                    "reply": reply,
                    "event_chip": {"city": nxt_city, "state": nxt_state,
                                   "dates": nxt.get("dates",""), "status": "next-upcoming",
                                   "confidence": "low"}}

    # Generic fallback
    return {"bucket": "B", "reason": "no KB pattern matched",
            "reply": None}




# ============================================================================
# 2026-07-02 — Claude API engine ("smart" classifier)
# Understands ANY phrasing (English + Spanish), replies in the commenter's
# language (EN/ES ONLY — NEVER Hebrew), grounded in the KB schedule + venues.
# Fail-safe: no ANTHROPIC_API_KEY / API error / invalid output → falls back to
# the keyword classify() below, so the agent NEVER goes dark.
# ============================================================================

_LLM_MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-haiku-4-5")
_HEBREW_RE = re.compile(r"[\u0590-\u05FF]")

# Spanish safety keywords — hard guardrails BEFORE the LLM (complaints must
# always reach Lauren, in any language)
_ES_ISSUE_RE = re.compile(
    r"\b(estafa|fraude|reembolso|queja|devoluci\u00f3n|devolucion|roto|da\u00f1ado|"
    r"danado|defectuoso|nunca (lleg\u00f3|llego|recib\u00ed|recibi)|no me (lleg\u00f3|llego)|"
    r"compr\u00e9|compre y|me enga\u00f1aron)\b", re.IGNORECASE)


def _llm_system_prompt(kb: dict) -> str:
    today = dt.date.today().strftime("%B %d, %Y")
    # 2026-07-02 — split into two explicit sections (inline PAST/UPCOMING tags
    # still produced tense mistakes: "we just hit San Jose Aug 28-30" for an
    # upcoming date)
    _up, _past = [], []
    for city, dates in kb.get("schedule", {}).items():
        (_past if _is_past(dates) else _up).append(f"- {city.title()}: {dates}")
    sched_lines = ("### UPCOMING events (these have NOT happened yet — always future tense):\n"
                   + "\n".join(_up)
                   + "\n\n### PAST 2026 events (already happened — always past tense, never invite):\n"
                   + "\n".join(_past))
    faq_lines = "\n".join(f"Q: {f['q']}\nA: {f['a']}" for f in kb.get("faqs", []))
    venues = kb.get("_venues", [])
    venue_lines = "\n".join(f"- {v['city']}: {v['venue']}, {v['address']} ({v['dates']})"
                            for v in venues[-25:])
    return f"""You are the social-media assistant for The Makeup Blowout Sale — a traveling weekend makeup sale (Fri-Sun, 10am-5pm, FREE admission, up to 75% off retail). Today is {today}.

You classify ONE Instagram/Facebook comment or DM and, when appropriate, draft the public reply.

## Business facts (the ONLY facts you may use)
- Hours: Friday, Saturday & Sunday, 10am-5pm. Free admission, no ticket needed. Free parking. Kids welcome.
- Payments: cash, all cards, tap-to-pay, Apple Pay.
- Sign-up gift: FREE Eyeshadow on arrival. Share-our-post bonus gift: FREE Glitter Liner.
- We ROTATE cities: each city gets a sale roughly once every 1-2 years. If a city is not on the schedule below, it is NOT happening in 2026 — never promise a city or date not listed.
- Website for online info: https://www.themakeupblowout.com
- FULL tour schedule page (all cities & dates): https://www.themakeupblowout.com/#events

## 2026 schedule
{sched_lines}

## Recent/upcoming venue addresses
{venue_lines}

## FAQ (use these answers, may rephrase)
{faq_lines}

## Post context
If a post caption is provided in the user message, the comment is about THAT event — use its city/dates/venue.

## Classification buckets
- "A"   = simple question you can answer confidently from the facts above (city/state/schedule, hours, admission, parking, payment, kids, brands, mystery box, where/address, online ordering). Auto-reply will be posted publicly.
- "B"   = needs the owner personally: complaints, refunds, past purchases, damaged/missing items, influencer/collab/PR/press/job/vendor inquiries, anything personal, anything you are not SURE about.
- "NEG" = hostile/skeptical (scam accusations, insults). Never reply.
- "SKIP"= no reply needed: friend tags, emoji-only, plain praise with no question.

## Reply rules (bucket A only)
1. Reply in the SAME language as the comment — English or Spanish ONLY. NEVER Hebrew, never any other language.
2. Warm, playful brand voice ("girl", "babe", emojis 💄✨💕) — match the energy. In Spanish use the same warmth (amiga, chica).
3. Max 300 characters. No hashtags. No links EXCEPT https://www.themakeupblowout.com/#events — DO include that link whenever someone asks generally where/when the events are, which cities you visit, or for the schedule (and it's a good add-on for city questions too).
4. NEVER invent cities, dates, venues, or discounts. Only the schedule above. City asked about but not listed → rotation policy: not this year, we rotate every 1-2 years, it goes on the 2027 wishlist, follow us for announcements. City/date already passed → "we were just there {{dates}}" + rotation message. State asked about → mention that state's upcoming events if any; if only past ones, say we already came this year.
5. If the event in the post context already ended and the comment is just excitement ("can't wait!"), use SKIP.

## Output — STRICT JSON only, no other text
{{"bucket":"A|B|NEG|SKIP","language":"en|es","confidence":0.0-1.0,"reason":"<short>","reply":"<text or null>"}}"""


def _llm_call(system: str, user_msg: str, api_key: str) -> dict:
    import urllib.request as _rq
    body = json.dumps({
        "model": _LLM_MODEL,
        "max_tokens": 500,
        "system": [{"type": "text", "text": system,
                     "cache_control": {"type": "ephemeral"}}],
        "messages": [{"role": "user", "content": user_msg}],
    }).encode("utf-8")
    req = _rq.Request("https://api.anthropic.com/v1/messages", data=body, method="POST",
                      headers={"Content-Type": "application/json",
                               "x-api-key": api_key,
                               "anthropic-version": "2023-06-01"})
    with _rq.urlopen(req, timeout=45) as r:
        out = json.loads(r.read().decode("utf-8"))
    return out


def classify_llm(text: str, kb: dict) -> dict:
    """Claude-based classification. Raises on any problem (caller falls back)."""
    api_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("no ANTHROPIC_API_KEY")
    post_ctx = kb.get("_post_context") or {}
    caption = (post_ctx.get("caption") or "").strip()
    user_msg = ""
    if caption:
        user_msg += f"POST CAPTION (context):\n{caption[:900]}\n\n"
    user_msg += f"COMMENT TO CLASSIFY:\n{text.strip()[:1000]}"
    # 2026-07-03 — retry w/ backoff: the ads backfill hit rate limits (98
    # fallbacks). Capture the real error body; back off on 429/400/529.
    import time as _time, urllib.error as _ue
    out = None; last = None
    for _att in range(4):
        try:
            out = _llm_call(_llm_system_prompt(kb), user_msg, api_key)
            break
        except _ue.HTTPError as e:
            body = ""
            try: body = e.read().decode()[:400]
            except Exception: pass
            last = f"HTTP {e.code}: {body}"
            # Only back off on TRANSIENT errors. A 400 (e.g. "credit balance too
            # low") is permanent — retrying 4x just hammers the API and slows the
            # whole run by minutes. Fail fast so the keyword fallback kicks in.
            if e.code in (429, 500, 529) and _att < 3:
                _time.sleep(3 * (_att + 1))
                continue
            raise RuntimeError(last)
    if out is None:
        raise RuntimeError(last or "no response")
    raw = "".join(b.get("text", "") for b in out.get("content", []) if b.get("type") == "text").strip()
    # tolerate markdown fences / trailing prose — extract the FIRST JSON object
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw)
    m_json = re.search(r"\{.*?\}(?=\s*$)|\{.*\}", raw, re.S)
    res = json.loads(m_json.group(0) if m_json else raw)
    bucket = res.get("bucket")
    if bucket not in ("A", "B", "NEG", "SKIP"):
        raise ValueError(f"bad bucket {bucket!r}")
    reply = (res.get("reply") or "").strip() or None
    conf = float(res.get("confidence") or 0)
    lang = res.get("language") or "en"
    # --- validation gates ---
    if bucket == "A":
        if not reply:
            raise ValueError("bucket A with no reply")
        if _HEBREW_RE.search(reply):
            raise ValueError("Hebrew leaked into reply")
        if len(reply) > 500:
            reply = reply[:500]
        if conf < 0.75:
            # not sure enough to auto-post — keep the draft but let Lauren review
            return {"bucket": "B", "engine": "llm",
                    "reason": f"LLM low-confidence ({conf:.2f}): {res.get('reason','')}"[:160],
                    "reply": reply}
    return {"bucket": bucket, "engine": "llm",
            "reason": f"LLM ({lang}, {conf:.2f}): {res.get('reason','')}"[:160],
            "reply": reply if bucket == "A" else (reply or None)}


# 2026-07-03 - LLM health tracker. The Claude engine silently 400'd on EVERY
# comment ("credit balance too low") for ~2 days, unnoticed because the keyword
# fallback kept working. main() SMSes Lauren (once/day) if the engine is down.
_LLM_HEALTH = {"attempts": 0, "ok": 0, "credit_low": False, "last_err": ""}


def classify_smart(text: str, kb: dict) -> dict:
    """Production entry point: deterministic guardrails → LLM → keyword fallback.

    The guardrails run FIRST and are never overridable by the LLM:
    negative/scam + customer-issue + influencer patterns always protect Lauren.
    """
    t = (text or "").strip()
    if not t or len(t) < 2:
        return {"bucket": "SKIP", "reason": "empty/too short (auto-done)", "reply": None}
    # hard guardrails (any language)
    if NEGATIVE_KEYWORDS.search(t):
        return {"bucket": "NEG", "reason": "negative keywords detected", "reply": None}
    if CUSTOMER_ISSUE_PAT.search(t) or _ES_ISSUE_RE.search(t):
        return {"bucket": "B", "reason": "Customer issue / past purchase — needs Lauren personally",
                "reply": None}
    if re.search(r"\b(content\s+creator|influencer|collab(?:oration)?|partnership|"
                 r"pr\s+list|brand\s+ambassador)\b", t, re.IGNORECASE):
        return {"bucket": "B", "reason": "Influencer/collab inquiry — Lauren personally", "reply": None}
    if _is_tag_only(t) and not _place_with_state(t):
        return {"bucket": "SKIP", "reason": "tag-only-friends (no question)", "reply": None}
    if _is_emoji_only_friendly(t) or _is_reaction_or_empty(t):
        return {"bucket": "SKIP", "reason": "emoji/reaction only (auto-done)", "reply": None}
    # LLM engine
    if os.environ.get("ANTHROPIC_API_KEY", "").strip():
        try:
            import time as _t
            _now = _t.time()
            _last = globals().get("_LLM_LAST_CALL", 0)
            if _now - _last < 1.3:            # ~45 req/min ceiling
                _t.sleep(1.3 - (_now - _last))
            globals()["_LLM_LAST_CALL"] = _t.time()
            _LLM_HEALTH["attempts"] += 1
            _res = classify_llm(t, kb)
            _LLM_HEALTH["ok"] += 1
            return _res
        except Exception as e:
            _LLM_HEALTH["attempts"] += 1
            _emsg = str(e)
            _LLM_HEALTH["last_err"] = _emsg[:200]
            if "credit balance" in _emsg.lower() or "plans & billing" in _emsg.lower():
                _LLM_HEALTH["credit_low"] = True
            print(f"  ⚠ LLM classify failed ({_emsg[:300]}) — keyword fallback")
    # keyword engine (fallback / no key)
    res = classify(t, kb)
    res.setdefault("engine", "keywords")
    return res


# --- Preview HTML builder ----------------------------------------------------

PAGE_HEAD = """<!DOCTYPE html>
<html lang="he" dir="rtl">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>תיבת מטה — The Makeup Blowout</title>
<style>
:root{--bg:#0f1117;--card:#171b26;--card2:#1d2331;--line:#262c3d;--txt:#e8eaf0;
 --dim:#9aa3b5;--pink:#f01070;--green:#22c55e;--amber:#f59e0b;--red:#ef4444;--blue:#3b82f6}
*{box-sizing:border-box}
body{background:var(--bg);color:var(--txt);font-family:-apple-system,"Segoe UI",Heebo,Arial,sans-serif;
 margin:0;padding:20px;max-width:840px;margin-inline:auto;line-height:1.55}
a{color:var(--blue)}
.topbar{display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px;margin-bottom:14px}
.back{background:var(--card);color:var(--dim);padding:8px 14px;border-radius:10px;text-decoration:none;font-size:13px;font-weight:600;border:1px solid var(--line)}
h1{font-size:24px;margin:4px 0 2px}
.fresh{display:inline-flex;gap:6px;align-items:center;background:var(--card);border:1px solid var(--line);
 color:var(--dim);border-radius:999px;padding:6px 14px;font-size:12.5px;font-weight:600}
.fresh .dot{width:8px;height:8px;border-radius:50%;background:var(--green);display:inline-block}
.fresh.stale .dot{background:var(--amber)}

/* HERO */
.hero{border-radius:16px;padding:22px 24px;margin:14px 0;display:flex;gap:16px;align-items:center}
.hero .big{font-size:42px;line-height:1}
.hero .t1{font-size:20px;font-weight:800;margin:0}
.hero .t2{font-size:13.5px;color:rgba(255,255,255,.75);margin-top:3px}
.hero-ok{background:linear-gradient(135deg,#0c3d24,#14532d);border:1px solid #1f7a44}
.hero-warn{background:linear-gradient(135deg,#402508,#713f12);border:1px solid #b45309}
.hero-neg{background:linear-gradient(135deg,#450a0a,#7f1d1d);border:1px solid #b91c1c}

/* STATS */
.stats{display:grid;grid-template-columns:repeat(4,1fr);gap:10px;margin:14px 0}
.stat{background:var(--card);border:1px solid var(--line);border-radius:14px;padding:14px 10px;text-align:center}
.stat .num{font-size:26px;font-weight:800}
.stat .lbl{font-size:11.5px;color:var(--dim);margin-top:3px;font-weight:600}
.stat.green .num{color:var(--green)} .stat.amber .num{color:var(--amber)}
.stat.gray .num{color:var(--dim)} .stat.blue .num{color:var(--blue)}
.coverage{background:var(--card);border:1px dashed var(--line);border-radius:12px;padding:10px 16px;
 font-size:12.5px;color:var(--dim);margin:10px 0 18px}

/* ATTENTION LIST */
.attention-block{margin:10px 0}
.attention-block h2{font-size:17px;margin:0 0 10px}
.attention-list{display:flex;flex-direction:column;gap:12px}
.attention-row{background:var(--card);border:1px solid var(--line);border-radius:14px;padding:14px 16px}
.attention-row.neg{border-color:#7f1d1d;background:#1c1214}
.att-meta{display:flex;flex-wrap:wrap;gap:6px;align-items:center;margin-bottom:8px}
.chip{display:inline-block;border-radius:999px;padding:3px 10px;font-size:11.5px;font-weight:700}
.chip.ig{background:#3a1030;color:#f9a8d4}.chip.fb{background:#102a4a;color:#93c5fd}
.chip.dm{background:#0f3730;color:#6ee7b7}
.chip.neg{background:#450a0a;color:#fca5a5}
.chip.carried{background:#3a2a0a;color:#fbbf24}
.att-event{font-size:11.5px;color:#c4b5fd;background:#241b3f;border-radius:999px;padding:3px 10px;font-weight:700}
.att-event.past{color:#fca5a5;background:#3f1b1b}
.att-time{font-size:11.5px;color:var(--dim);margin-inline-start:auto}
.att-time.old{color:var(--amber);font-weight:700}
.who{font-weight:800;font-size:14px}
.att-message{background:var(--card2);border-radius:10px;padding:10px 12px;font-size:14px;white-space:pre-wrap;word-break:break-word}
.att-why{font-size:12px;color:var(--dim);margin-top:6px}
textarea.att-reply{width:100%;background:#12161f;color:var(--txt);border:1px solid var(--line);
 border-radius:10px;padding:10px 12px;font-family:inherit;font-size:13.5px;min-height:64px;margin-top:10px;direction:ltr;text-align:left}
.att-actions{display:flex;gap:8px;flex-wrap:wrap;margin-top:8px}
.send-btn{background:var(--pink);color:#fff;border:none;border-radius:9px;padding:9px 16px;font-weight:800;
 font-size:13px;cursor:pointer;font-family:inherit}
.send-btn:hover{filter:brightness(1.12)} .send-btn[disabled]{opacity:.55}
.send-btn.sent{background:#15803d}
.done-btn{background:transparent;color:var(--green);border:1px solid #1f7a44;border-radius:9px;padding:9px 16px;
 font-weight:700;font-size:13px;cursor:pointer;font-family:inherit}
.done-btn:hover{background:#0c3d24} .done-btn[disabled]{opacity:.5}
.reply-btn-secondary{color:var(--dim);border:1px solid var(--line);border-radius:9px;padding:9px 14px;
 font-size:13px;text-decoration:none;font-weight:600}
.handled-fade{opacity:.25;transition:opacity .4s}
/* comment moderation (hide/delete/block) — injected by runtime JS on comment rows */
.mod-actions{display:flex;gap:8px;flex-wrap:wrap;margin-top:10px;padding-top:10px;border-top:1px dashed var(--line)}
.mod-btn{background:transparent;border:1px solid var(--line);border-radius:9px;padding:7px 13px;
 font-size:12.5px;font-weight:700;cursor:pointer;font-family:inherit;color:var(--dim)}
.mod-btn:hover{filter:brightness(1.15)} .mod-btn[disabled]{opacity:.5}
.mod-btn.mod-hide{color:#c7a94a;border-color:#7a5f1f}
.mod-btn.mod-hide:hover{background:#2a2109}
.mod-btn.mod-del{color:#e06666;border-color:#7f2d2d}
.mod-btn.mod-del:hover{background:#2a1414}
.mod-btn.mod-block{color:#d17ad1;border-color:#6d2d6d}
.mod-btn.mod-block:hover{background:#241024}
.mod-btn.sent{background:#15803d;color:#fff;border-color:#15803d}
.mod-hidden-kept{opacity:.72;border-inline-start:3px solid #7a5f1f}
.mod-hidden-kept .att-message{text-decoration:line-through;color:var(--dim)}

/* TRANSPARENCY LOG */
details.log{background:var(--card);border:1px solid var(--line);border-radius:14px;margin:18px 0;overflow:hidden}
details.log summary{padding:14px 18px;cursor:pointer;font-weight:800;font-size:14.5px;list-style:none;color:var(--txt)}
details.log summary::-webkit-details-marker{display:none}
details.log summary::after{content:" ▼";color:var(--dim);font-size:11px}
details.log[open] summary::after{content:" ▲"}
.log-item{border-top:1px solid var(--line);padding:12px 18px;font-size:13px}
.log-item .q{color:var(--dim)} .log-item .a{background:#0c2a1a;border-radius:8px;padding:8px 10px;margin-top:6px;
 color:#bbf7d0;direction:ltr;text-align:left;white-space:pre-wrap;word-break:break-word}
.log-item .lwho{font-weight:700}
.log-item .lmeta{font-size:11px;color:var(--dim);margin-top:4px}

.legend{background:var(--card);border-radius:12px;border:1px solid var(--line);padding:12px 16px;font-size:12.5px;color:var(--dim);margin:14px 0}
.errors{background:#3a2a0a;border-radius:10px;padding:10px 16px;font-size:12px;color:#fbbf24;margin:10px 0}
.token-banner{background:#3a2a0a;border:1px solid #d97706;color:#fbbf24;padding:12px 16px;border-radius:10px;margin:12px 0;font-size:13px}
.token-banner input{background:#12161f;color:#eee;border:1px solid #555;padding:6px 10px;border-radius:6px;
 font-family:monospace;font-size:11px;width:100%;max-width:400px;margin-top:6px}
.token-banner button{background:#fbbf24;color:#1a1a2a;font-weight:700;padding:6px 14px;border-radius:6px;border:none;cursor:pointer;margin-inline-end:6px;font-family:inherit}
footer{color:#555;font-size:11.5px;margin:26px 0 8px;text-align:center}
@media (max-width:600px){
 body{padding:12px}
 .stats{grid-template-columns:repeat(2,1fr)}
 .hero{padding:16px}.hero .big{font-size:34px}.hero .t1{font-size:17px}
 .att-actions{flex-direction:column}.att-actions>*{width:100%;text-align:center}
}
</style>
</head>
<body>
"""

PAGE_FOOT = """
<footer>הדף מתעדכן אוטומטית בכל ריצה (כל 3 שעות) · <a href="../" style="color:#667">דשבורד @meta הישן</a></footer>
<script>
__RUNTIME_JS__
</script>
</body>
</html>"""



def _format_la_time(iso_str: str) -> str:
    """Convert ISO UTC string to LA-time formatted display.
    
    Input:  '2026-05-08T16:44:25Z' or '2026-05-08T16:44:25.123Z'
    Output: 'May 8, 2026 · 9:44 AM PDT' (or PST in winter)
    """
    import datetime as _dt
    try:
        s = iso_str.rstrip("Z").split(".")[0]
        utc_dt = _dt.datetime.fromisoformat(s).replace(tzinfo=_dt.timezone.utc)
        try:
            from zoneinfo import ZoneInfo
            la_dt = utc_dt.astimezone(ZoneInfo("America/Los_Angeles"))
        except Exception:
            la_dt = utc_dt.astimezone(_dt.timezone(_dt.timedelta(hours=-7)))  # fallback PDT
        # Format like "May 8, 2026 · 9:44 AM PDT"
        try:
            tz_name = la_dt.tzname() or "PT"
        except Exception:
            tz_name = "PT"
        return la_dt.strftime("%b %d, %Y · %-I:%M %p ") + tz_name
    except Exception:
        return iso_str  # fallback to raw

def _format_received(iso_str: str) -> tuple:
    """Format a received-timestamp for the per-item display.

    Returns (display_text, freshness_class) where freshness_class is one of
    "fresh" (< 24h), "" (1-3 days), or "old" (> 3 days).

    display_text is e.g. "May 16 · לפני 4 ימים" — short, LTR-friendly,
    with a Hebrew relative-time tail.

    Used by render_preview() to surface WHEN each Bucket-B item arrived,
    so Lauren can prioritize the freshest threads (where the customer is
    likely still waiting on the same screen).
    """
    import datetime as _dt
    if not iso_str:
        return ("", "")
    try:
        s = iso_str.rstrip("Z").split(".")[0].split("+")[0]
        utc_dt = _dt.datetime.fromisoformat(s).replace(tzinfo=_dt.timezone.utc)
        try:
            from zoneinfo import ZoneInfo
            la_dt = utc_dt.astimezone(ZoneInfo("America/Los_Angeles"))
        except Exception:
            la_dt = utc_dt.astimezone(_dt.timezone(_dt.timedelta(hours=-7)))
        now = _dt.datetime.now(_dt.timezone.utc)
        delta = now - utc_dt
        secs = int(delta.total_seconds())
        if secs < 0:
            secs = 0
        # Hebrew relative phrase
        if   secs < 60:        rel = "ממש עכשיו"
        elif secs < 3600:      rel = f"לפני {secs//60} דק'"
        elif secs < 86400:
            h = secs // 3600
            rel = "לפני שעה" if h == 1 else f"לפני {h} שעות"
        else:
            d = secs // 86400
            rel = "לפני יום" if d == 1 else f"לפני {d} ימים"
        absolute = la_dt.strftime("%b %-d · %-I:%M %p")
        if secs < 86400:
            cls = "fresh"
        elif secs < 3 * 86400:
            cls = ""
        else:
            cls = "old"
        return (f"{absolute} · {rel}", cls)
    except Exception:
        return (iso_str[:16], "")

def _make_draft_reply(name: str, customer_text: str, channel: str) -> str:
    """Generate a friendly placeholder draft for Lauren to edit before sending."""
    display_name = (name or "").lstrip("@")
    first_word = display_name.split()[0].split("_")[0] if display_name else "girl"
    first_word = first_word[:20].title() if first_word else "girl"

    text_lower = (customer_text or "").lower()

    if "?" in (customer_text or "") or any(w in text_lower for w in ["when","where","how","what","why","who","do you","can i","is it"]):
        return f"Hey {first_word} 💜 Thanks for reaching out! Let me check on that and get back to you super soon ✨"
    if any(w in text_lower for w in ["bad","terrible","wrong","broken","upset","angry","disappointed","scam","refund","return"]):
        return f"Hi {first_word} — so sorry to hear that 💜 Can you share a bit more so we can make this right?"
    if not (customer_text or "").strip() or any(emj in customer_text for emj in ["📸","🎁","❤️","🎬","🎤","📎"]):
        return f"Hi {first_word} 💜 Thanks so much for the message! Anything we can help with?"
    return f"Hi {first_word} 💜 Thanks for reaching out! Let me get back to you super soon ✨"


def render_preview(snapshot: dict, classified_messenger: list,
                   classified_fb_comments: list,
                   classified_ig_comments: list) -> str:
    """Build the static HTML preview page."""
    n_a = sum(1 for c in classified_messenger if c["cls"]["bucket"] == "A")
    n_b = sum(1 for c in classified_messenger if c["cls"]["bucket"] == "B")
    n_neg = sum(1 for c in classified_messenger if c["cls"]["bucket"] == "NEG")
    n_skip = sum(1 for c in classified_messenger if c["cls"]["bucket"] == "SKIP")
    n_fb = len(classified_fb_comments)
    n_ig = len(classified_ig_comments)

    # Collect every Bucket B item across all channels — the things needing Lauren.
    # Server-side filter: skip items already marked handled (so count + visible items match).
    try:
        from pathlib import Path as _Path
        import json as _json
        _h = _Path("docs/meta/handled.json")
        _handled = _json.loads(_h.read_text()) if _h.exists() else {}
    except Exception:
        _handled = {}
    def _is_handled(dkey):
        if not dkey: return False
        return bool(_handled.get(dkey, {}).get("handled"))

    attention_items = []
    for m in classified_messenger:
        if _is_handled(m.get("dedup_key", "")):
            continue
        if m["cls"]["bucket"] in ("B", "NEG"):
            who = m.get("name", "?")
            attention_items.append({
                "channel": "Messenger",
                "who":     who,
                "what":    m.get("msg", "(no text)") or "(empty)",
                "url":     m.get("reply_url", ""),
                "bucket":  m["cls"]["bucket"],
                "dedup_key": m.get("dedup_key", ""),
                "reply_kind": "messenger",
                "target_id": m.get("customer_psid", ""),  # PSID for Send API
                "received": m.get("updated_time", ""),
                "event_chip": m["cls"].get("event_chip"),
                "draft":    m["cls"].get("reply") or _make_draft_reply(who, m.get("msg", ""), "messenger"),
            })
    for c in classified_fb_comments:
        if _is_handled(c.get("dedup_key", "")):
            continue
        if c["cls"]["bucket"] in ("B", "NEG"):
            who_from = c.get("from", "?")
            attention_items.append({
                "channel": "FB comment",
                "who":     who_from,
                "what":    c.get("text", "(no text)"),
                "url":     c.get("reply_url", ""),
                "bucket":  c["cls"]["bucket"],
                "dedup_key": c.get("dedup_key", ""),
                "reply_kind": "fb_comment",
                "target_id": c.get("comment_id") or c.get("id", ""),
                "received": c.get("created_time", ""),
                "event_chip": c["cls"].get("event_chip"),
                "draft":    c["cls"].get("reply") or _make_draft_reply(who_from, c.get("text", ""), "fb_comment"),
            })
    for c in classified_ig_comments:
        if _is_handled(c.get("dedup_key", "")):
            continue
        if c["cls"]["bucket"] in ("B", "NEG"):
            who_ig = "@" + c.get("username", "?")
            attention_items.append({
                "channel": "IG comment",
                "who":     who_ig,
                "what":    c.get("text", "(no text)"),
                "url":     c.get("reply_url", ""),
                "bucket":  c["cls"]["bucket"],
                "dedup_key": c.get("dedup_key", ""),
                "reply_kind": "ig_comment",
                "target_id": c.get("id", ""),
                "received": c.get("timestamp", ""),
                "event_chip": c["cls"].get("event_chip"),
                "draft":    c["cls"].get("reply") or _make_draft_reply(who_ig, c.get("text", ""), "ig_comment"),
            })

    # ================================================================
    # 2026-07-02 — PERSISTENT PENDING QUEUE ("nothing falls through")
    # Items needing Lauren are merged into pending.json and carried across
    # runs until handled — even after the post leaves the scan window.
    # ================================================================
    pend_path = _Path("docs/meta/inbox-api-preview/pending.json")
    try:
        _old_pend = _json.loads(pend_path.read_text(encoding="utf-8")).get("items", {})
    except Exception:
        _old_pend = {}
    _now = _now_iso()
    cur_keys = set()
    for it in attention_items:
        k = it.get("dedup_key", "")
        cur_keys.add(k)
        prev = _old_pend.get(k) or {}
        it["first_seen"] = prev.get("first_seen") or it.get("received") or _now
    # 2026-07-03 — SELF-HEALING carry-purge. A comment that was RE-FETCHED this
    # run and is now classified SKIP (friend-tag / emoji / plain praise, no
    # question) must DROP OUT of the pending queue — NOT be re-carried as its
    # stale B entry. Root cause of the 2026-07-02 backfill pileup: the first
    # backfill hit LLM rate-limits (98 fallbacks) and the keyword fallback dumped
    # ~135 social-noise comments into Bucket B; without this purge the carry loop
    # re-added every one of them forever, even after the (now-fixed) LLM would
    # correctly SKIP them. We only purge SKIP (definitively no-reply-needed);
    # Bucket A that got auto-replied is already removed via handled.json.
    _resolved_skip = set()
    for _lst in (classified_messenger, classified_fb_comments, classified_ig_comments):
        for _c in _lst:
            if (_c.get("cls") or {}).get("bucket") == "SKIP":
                _rk = _c.get("dedup_key", "")
                if _rk:
                    _resolved_skip.add(_rk)
    for k, v in _old_pend.items():
        if k in cur_keys or _is_handled(k) or k in _resolved_skip:
            continue
        v["carried"] = True                     # out of scan window but still open
        attention_items.append(v)
    try:
        pend_path.parent.mkdir(parents=True, exist_ok=True)
        pend_path.write_text(_json.dumps(
            {"_updated_at": _now,
             "items": {it["dedup_key"]: it for it in attention_items if it.get("dedup_key")}},
            ensure_ascii=False, indent=1), encoding="utf-8")
    except Exception as _pe:
        print(f"  ⚠ pending.json write failed: {_pe}")
    # NEG first, then newest first
    attention_items.sort(key=lambda it: (it.get("bucket") != "NEG",
                                         it.get("received") or it.get("first_seen") or ""),
                         )
    attention_items.sort(key=lambda it: it.get("received") or it.get("first_seen") or "", reverse=True)
    attention_items.sort(key=lambda it: it.get("bucket") != "NEG")
    n_pending = len(attention_items)
    n_neg_all = sum(1 for it in attention_items if it.get("bucket") == "NEG")

    # counts across ALL channels
    def _cnt(lst, b): return sum(1 for c in lst if c["cls"]["bucket"] == b)
    auto_replied = (_cnt(classified_messenger, "A") + _cnt(classified_fb_comments, "A")
                    + _cnt(classified_ig_comments, "A"))
    skipped = (_cnt(classified_messenger, "SKIP") + _cnt(classified_fb_comments, "SKIP")
               + _cnt(classified_ig_comments, "SKIP"))
    total_comments = len(classified_fb_comments) + len(classified_ig_comments)
    total_dms = len(classified_messenger)

    _CH_CHIP = {"Messenger": '<span class="chip dm">💬 מסנג׳ר</span>',
                "FB comment": '<span class="chip fb">📘 פייסבוק</span>',
                "IG comment": '<span class="chip ig">📸 אינסטגרם</span>'}

    parts = [PAGE_HEAD]
    parts.append('<div class="topbar">'
                 '<div><h1>📬 תיבת מטה</h1>'
                 '<div style="font-size:12.5px;color:var(--dim)">מסנג׳ר · תגובות פייסבוק · תגובות אינסטגרם — מנוהל אוטומטית</div></div>'
                 f'<div><span class="fresh" id="fresh" data-fetched="{html.escape(snapshot["fetched_at"])}">'
                 '<span class="dot"></span><span id="fresh-txt">…</span></span>'
                 '<div style="text-align:center;font-size:11px;color:var(--dim);margin-top:4px" id="next-run"></div></div>'
                 '<a class="back" href="https://dashboard.themakeupblowout.com/">← Agent Hub</a>'
                 '</div>')

    # === HERO ===
    if n_pending == 0:
        parts.append('<div class="hero hero-ok" id="hero">'
                     '<div class="big">✅</div><div>'
                     '<p class="t1">הכל מטופל — שום תגובה לא מחכה לך</p>'
                     '<p class="t2">כל שאלה שנסרקה נענתה אוטומטית או סומנה כטופלה. פריטים פתוחים נשמרים כאן גם בין ריצות — כלום לא הולך לאיבוד.</p>'
                     '</div></div>')
    elif n_neg_all:
        parts.append(f'<div class="hero hero-neg" id="hero">'
                     f'<div class="big">🔴</div><div>'
                     f'<p class="t1"><span id="pending-count">{n_pending}</span> ממתינות לך — כולל {n_neg_all} רגישות</p>'
                     f'<p class="t2">תגובות רגישות (תלונה/ספקנות) אף פעם לא נענות אוטומטית — רק את. השאר למטה עם טיוטה מוכנה.</p>'
                     f'</div></div>')
    else:
        parts.append(f'<div class="hero hero-warn" id="hero">'
                     f'<div class="big">👀</div><div>'
                     f'<p class="t1"><span id="pending-count">{n_pending}</span> תגובות מחכות להחלטה שלך</p>'
                     f'<p class="t2">לכל אחת יש טיוטה מוכנה — עריכה (אם צריך) → שליחה בלחיצה. הפריטים נשמרים בין ריצות עד שמטפלים בהם.</p>'
                     f'</div></div>')

    # === STATS ===
    parts.append('<div class="stats">'
                 f'<div class="stat green"><div class="num">{auto_replied}</div><div class="lbl">🤖 נענו אוטומטית (בסריקה הזו)</div></div>'
                 f'<div class="stat amber"><div class="num" id="stat-pending">{n_pending}</div><div class="lbl">👀 מחכות לך</div></div>'
                 f'<div class="stat gray"><div class="num">{skipped}</div><div class="lbl">😴 לא דורשות מענה</div></div>'
                 f'<div class="stat blue"><div class="num">{total_comments + total_dms}</div><div class="lbl">📡 נבדקו בסריקה</div></div>'
                 '</div>')
    parts.append(f'<div class="coverage">📡 <b>מה כוסה בסריקה האחרונה:</b> 25 הפוסטים האחרונים בפייסבוק + 25 באינסטגרם (כל התגובות, כולל ישנות) + {total_dms} שיחות מסנג׳ר · רץ אוטומטית כל 3 שעות · שאלות פשוטות נענות מיד ע״י מנוע Claude (אנגלית + ספרדית) · תלונות ופניות אישיות תמיד מחכות לך.</div>')

    # === ATTENTION LIST ===
    if attention_items:
        parts.append('<div class="attention-block">')
        parts.append(f'<h2 id="att-title">👀 ממתינות לטיפול שלך ({n_pending})</h2>')
        parts.append('<div class="attention-list">')
        for item in attention_items:
            url_attr = html.escape(item.get("url") or "") or "#"
            dkey = html.escape(item.get("dedup_key", ""))
            tid  = html.escape(item.get("target_id", ""))
            kind = html.escape(item.get("reply_kind", ""))
            draft = html.escape(item.get("draft") or "")
            is_neg = item.get("bucket") == "NEG"
            row_cls = "attention-row neg" if is_neg else "attention-row"
            parts.append(f'<div class="{row_cls}" data-dedup-key="{dkey}" data-target-id="{tid}" data-reply-kind="{kind}">')
            chips = [_CH_CHIP.get(item.get("channel", ""), "")]
            if is_neg:
                chips.append('<span class="chip neg">⚠️ רגיש — בלי מענה אוטומטי</span>')
            if item.get("carried"):
                chips.append('<span class="chip carried">⏳ ממתין מריצה קודמת</span>')
            ev = item.get("event_chip") or {}
            if ev.get("city"):
                ev_cls = "past" if ev.get("status") == "past" else ""
                chips.append(f'<span class="att-event {ev_cls}">🎯 {html.escape(ev.get("city",""))}, {html.escape(ev.get("state",""))} · {html.escape((ev.get("dates","") or "")[:24])}</span>')
            recv_text, recv_cls = _format_received(item.get("received", "") or item.get("first_seen", ""))
            time_html = (f'<span class="att-time {recv_cls}">🕒 {html.escape(recv_text)}</span>' if recv_text else "")
            parts.append(f'<div class="att-meta">{"".join(chips)}<span class="who">{html.escape(item.get("who","?"))}</span>{time_html}</div>')
            parts.append(f'<div class="att-message">{html.escape(item.get("what") or "(ללא טקסט)")}</div>')
            parts.append(f'<textarea class="att-reply" data-key="{dkey}" placeholder="כתבי תשובה (אנגלית/ספרדית)...">{draft}</textarea>')
            parts.append('<div class="att-actions">')
            if tid and kind:
                parts.append(f'<button class="send-btn" data-key="{dkey}" data-tid="{tid}" data-kind="{kind}" onclick="sendReply(this)">📤 שלחי את התשובה</button>')
            if dkey:
                parts.append(f'<button class="done-btn" data-key="{dkey}" onclick="markDone(this)">✓ סמני כטופל (בלי לענות)</button>')
            if item.get("url"):
                parts.append(f'<a class="reply-btn-secondary" href="{url_attr}" target="_blank" rel="noopener" onclick="markReplyClicked(this)" data-key="{dkey}">↗ פתחי במטה</a>')
            parts.append('</div></div>')
        parts.append('</div></div>')

    # === TRANSPARENCY LOG — what the agent answered this scan ===
    replied = []
    for c in classified_messenger:
        if c["cls"]["bucket"] == "A" and c["cls"].get("reply"):
            replied.append(("💬", c.get("name","?"), c.get("msg",""), c["cls"]["reply"], c.get("reply_url",""), c["cls"].get("reason","")))
    for c in classified_fb_comments:
        if c["cls"]["bucket"] == "A" and c["cls"].get("reply"):
            replied.append(("📘", c.get("from","?"), c.get("text",""), c["cls"]["reply"], c.get("reply_url",""), c["cls"].get("reason","")))
    for c in classified_ig_comments:
        if c["cls"]["bucket"] == "A" and c["cls"].get("reply"):
            replied.append(("📸", "@"+c.get("username","?"), c.get("text",""), c["cls"]["reply"], c.get("reply_url",""), c["cls"].get("reason","")))
    if replied:
        parts.append('<details class="log">')
        parts.append(f'<summary>🤖 מה הסוכן ענה בסריקה הזו ({len(replied)}) — שקיפות מלאה</summary>')
        for icon, who, q, a, url, reason in replied:
            parts.append('<div class="log-item">')
            parts.append(f'<div class="lwho">{icon} {html.escape(who)}</div>')
            parts.append(f'<div class="q">שאלה: {html.escape(q[:220])}</div>')
            parts.append(f'<div class="a">{html.escape(a)}</div>')
            link = f' · <a href="{html.escape(url)}" target="_blank" rel="noopener">פתחי</a>' if url else ""
            parts.append(f'<div class="lmeta">{html.escape(reason[:110])}{link}</div>')
            parts.append('</div>')
        parts.append('</details>')

    if snapshot.get("errors"):
        parts.append('<div class="errors"><b>שגיאות לא-קריטיות בסריקה:</b> ' +
                     " · ".join(html.escape(e[:120]) for e in snapshot["errors"][:5]) + '</div>')

    parts.append('<div class="legend">ℹ️ <b>איך זה עובד:</b> הסוכן סורק כל 3 שעות. שאלות פשוטות (לו״ז ערים, שעות, חניה, תשלום, מותגים — בכל ניסוח, גם ספרדית) נענות אוטומטית ומופיעות בלוג השקיפות. כל השאר מחכה לך כאן עם טיוטה — וגם אם הפוסט יוצא מחלון הסריקה, הפריט נשאר עד שתטפלי בו. תגובה שנעלמה מהרשימה = טופלה.</div>')

    # freshness + next-run clock
    parts.append("""<script>
(function(){
  var el=document.getElementById('fresh'),tx=document.getElementById('fresh-txt'),nr=document.getElementById('next-run');
  if(!el)return;
  var fetched=new Date(el.dataset.fetched);
  function tick(){
    var m=Math.max(0,Math.round((Date.now()-fetched.getTime())/60000));
    tx.textContent = m<60 ? ('עודכן לפני '+m+' דק׳') : ('עודכן לפני '+Math.round(m/60)+' שע׳');
    if(m>200){el.classList.add('stale');}
    var now=new Date(); var nx=new Date(now);
    nx.setUTCMinutes(0,0,0); nx.setUTCHours(now.getUTCMinutes()||now.getUTCSeconds()?now.getUTCHours()+1:now.getUTCHours());
    while(nx.getUTCHours()%3!==0){nx.setUTCHours(nx.getUTCHours()+1);}
    if(nx<=now){nx.setUTCHours(nx.getUTCHours()+3);}
    nr.textContent='סריקה הבאה ~'+nx.toLocaleTimeString('he-IL',{hour:'2-digit',minute:'2-digit'});
  }
  tick(); setInterval(tick,30000);
})();
window.__updatePendingCount=function(n){
  var pc=document.getElementById('pending-count'); if(pc)pc.textContent=n;
  var sp=document.getElementById('stat-pending'); if(sp)sp.textContent=n;
  var at=document.getElementById('att-title'); if(at)at.textContent='👀 ממתינות לטיפול שלך ('+n+')';
  if(n===0){var h=document.getElementById('hero'); if(h){h.className='hero hero-ok';
    h.innerHTML='<div class="big">✅</div><div><p class="t1">הכל מטופל — שום תגובה לא מחכה לך</p><p class="t2">כל הכבוד! הרשימה נקייה.</p></div>';}}
};
</script>""")

    rendered = "\n".join(parts) + PAGE_FOOT
    # Inject the runtime JS (kept in a separate file to avoid Python escape hell)
    js_path = Path(__file__).resolve().parent / "meta_inbox_preview_runtime.js"
    runtime_js = js_path.read_text(encoding="utf-8") if js_path.exists() else ""
    return rendered.replace("__RUNTIME_JS__", runtime_js)


# --- Main --------------------------------------------------------------------

def _resolve_kb_path() -> Path:
    """Find the KB docx — checked in to the repo, with Cowork fallback."""
    candidates = [
        Path(__file__).resolve().parent / "data" / "meta_inbox_kb.docx",
        Path("/sessions/dreamy-compassionate-wozniak/mnt/Claude/Scheduled/NEW/meta-inbox/inbox_knowledge_base.docx"),
    ]
    for c in candidates:
        if c.exists():
            return c
    raise FileNotFoundError(f"meta_inbox_kb.docx not found in any of: {candidates}")


def _now_iso() -> str:
    return dt.datetime.now(dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def main():
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("--send-sms", action="store_true",
                    help="Send Touchpoint 1 + Touchpoint 3 SMS to Lauren")
    ap.add_argument("--reply-test", metavar="CONV_ID",
                    help="Phase 2 supervised test: send the Bucket A reply for ONE specific Messenger conv_id (live)")
    ap.add_argument("--reply-bucket-a", action="store_true",
                    help="Phase 2 LIVE: actually send all Bucket A Messenger replies (dry_run=False). Use only after a successful --reply-test.")
    ap.add_argument("--bulk-mark-handled", action="store_true",
                    help="Mark ALL classified items (A+B+NEG) as handled WITHOUT sending — for manual cleanup by Lauren.")
    args = ap.parse_args()

    repo = Path(__file__).resolve().parent.parent
    kb_path = _resolve_kb_path()
    kb = load_kb(kb_path)
    kb["_venues"] = load_venues()
    print(f"KB loaded: {len(kb['faqs'])} FAQs, {len(kb['schedule'])} cities, "
          f"{len(kb['negatives'])} negative situations")

    print("Fetching inbox snapshot...")
    snap = fetch_recent_inbox(days=7, include_messenger=True, include_ig_dms=False,
                              include_fb_comments=True, include_ig_comments=True,
                              fb_post_limit=25, ig_media_limit=25)
    print(f"  messenger: {len(snap['messenger'])}, "
          f"fb_comments groups: {len(snap['fb_comments'])}, "
          f"ig_comments groups: {len(snap['ig_comments'])}")

    # 2026-07-02 — AD comments (Lauren: "לא לפספס שום מודעה או תגובה בשום מקום").
    # Dark ad posts don't appear in /{page}/posts; pull them via the Marketing API.
    from lauren_meta import fetch_ads_comments
    _ads_since = os.environ.get("ADS_BACKFILL_SINCE", "").strip() or None
    try:
        ad_fb, ad_ig, ad_errs = fetch_ads_comments(all_since=_ads_since)
        known_fb = {g["post"].get("id") for g in snap["fb_comments"]}
        known_ig = {g["media"].get("id") for g in snap["ig_comments"]}
        snap["fb_comments"] += [g for g in ad_fb if g["post"].get("id") not in known_fb]
        snap["ig_comments"] += [g for g in ad_ig if g["media"].get("id") not in known_ig]
        snap["errors"] += ad_errs[:10]
        print(f"  ads coverage ({'ALL since ' + _ads_since if _ads_since else 'ACTIVE only'}): "
              f"+{len(ad_fb)} fb ad-posts, +{len(ad_ig)} ig ad-media with comments"
              f" ({len(ad_errs)} errors)")
    except Exception as e:
        print(f"  ⚠ ads coverage failed (non-fatal): {e}")

    # Load dedup memory + filter items Lauren has already handled
    handled = load_handled()
    print(f"  handled-state entries: {len(handled)}")

    # Enrich messenger conversations with the latest message text
    classified_messenger = []
    for c in snap["messenger"]:
        if is_handled(handled, "messenger", c.get("id", "")):
            continue   # Lauren replied / marked done in a prior run
        conv_id = c["id"]
        # Get participants (filter out the Page itself)
        parts = c.get("participants", {}).get("data", [])
        customer = next((p for p in parts
                         if p.get("name") != "The Makeup Blowout Sale Group"), parts[0] if parts else {})
        # Fetch the most recent message text
        try:
            messages = fetch_messenger_messages(conv_id, limit=3)
            last_customer_msg = next((m for m in messages
                                      if m.get("from", {}).get("name") != "The Makeup Blowout Sale Group"),
                                     None)
            if last_customer_msg:
                text = last_customer_msg.get("message", "") or ""
                # If text is empty, derive a label from attachments/stickers
                if not text.strip():
                    if last_customer_msg.get("sticker"):
                        text = "🎁 (sticker)"
                    else:
                        atts = (last_customer_msg.get("attachments") or {}).get("data") or []
                        if atts:
                            att = atts[0]
                            mime = (att.get("mime_type") or "").lower()
                            atype = (att.get("type") or "").lower()
                            if "image" in mime or atype == "image":
                                text = "📸 (image attachment)"
                            elif "video" in mime or atype == "video":
                                text = "🎬 (video attachment)"
                            elif atype == "audio":
                                text = "🎤 (audio message)"
                            elif atype == "file":
                                text = "📎 (file attachment)"
                            else:
                                text = f"📎 ({atype or mime or 'attachment'})"
                        else:
                            text = "❤️ (reaction or empty message)"
            else:
                text = ""
        except Exception as e:
            text = f"(error fetching: {e})"
        kb["_seed"] = conv_id  # per-conversation seed for reply variation
        kb["_post_context"] = None  # DMs have no post context
        cls = classify_smart(text, kb)
        # Build a direct reply URL using the customer's PSID (the participant
        # who is NOT the Page). Format `https://www.facebook.com/messages/t/{psid}`
        # is the only reliable deep-link that opens the specific Messenger thread.
        # The conv_id (e.g. "t_10242117887745746") is NOT a usable URL component.
        page_id = get_fb_page_id()
        customer_psid = ""
        for pp in (c.get("participants", {}).get("data", []) or []):
            if pp.get("id") and pp.get("id") != page_id:
                customer_psid = pp["id"]
                break
        reply_url = f"https://www.facebook.com/messages/t/{customer_psid}" if customer_psid else \
                    f"https://business.facebook.com/latest/inbox/all?asset_id={page_id}"
        # 2026-05-14 — sentiment + urgency tags
        sentiment = classify_sentiment(text)
        urgent = is_urgent(text)
        classified_messenger.append({
            "conv_id": conv_id,
            "dedup_key": dedup_key("messenger", conv_id),
            "name": customer.get("name", "?"),
            "msg": text,
            "updated_time": c.get("updated_time", ""),
            "reply_url": reply_url,
            "customer_psid": customer_psid,
            "cls": cls,
            "sentiment": sentiment,
            "urgent": urgent,
        })

    # FB comments
    classified_fb = []
    for grp in snap["fb_comments"]:
        post_msg = grp["post"].get("message", "")[:80]
        post_permalink = grp["post"].get("permalink_url", "")
        for cmt in grp["comments"]:
            if is_handled(handled, "fb-comment", cmt.get("id", "")):
                continue
            # 2026-07-02 — already answered publicly (by Lauren or the agent) → skip
            _replies = (cmt.get("comments") or {}).get("data", [])
            if any((r.get("from") or {}).get("id") == get_fb_page_id() for r in _replies):
                continue
            txt = cmt.get("message", "")
            kb["_seed"] = cmt.get("id", "")
            kb["_post_context"] = {
                "caption": grp["post"].get("message", ""),
                "date": grp["post"].get("created_time", ""),
            }
            cls = classify_smart(txt, kb)
            # FB comment: prefer post permalink (the comment will be visible there;
            # Meta doesn't always expose direct comment-permalinks for replies)
            classified_fb.append({
                "comment_id": cmt.get("id", ""),
                "dedup_key": dedup_key("fb-comment", cmt.get("id", "")),
                "post_msg": post_msg,
                "from": cmt.get("from", {}).get("name", "?"),
                "text": txt,
                "created_time": cmt.get("created_time", ""),
                "reply_url": post_permalink or f"https://www.facebook.com/{cmt.get('id','')}",
                "cls": cls,
                "sentiment": classify_sentiment(txt),
                "urgent": is_urgent(txt),
            })

    # IG comments
    classified_ig = []
    for grp in snap["ig_comments"]:
        media_caption = (grp["media"].get("caption") or "").split("\n")[0][:80]
        media_permalink = grp["media"].get("permalink", "")
        for cmt in grp["comments"]:
            if is_handled(handled, "ig-comment", cmt.get("id", "")):
                continue
            # 2026-07-02 — already answered publicly → skip
            _reps = (cmt.get("replies") or {}).get("data", [])
            if any(r.get("username") == "themakeupblowoutsale" for r in _reps):
                continue
            txt = cmt.get("text", "")
            kb["_seed"] = cmt.get("id", "")
            kb["_post_context"] = {
                "caption": grp["media"].get("caption", ""),
                "date": grp["media"].get("timestamp", ""),
            }
            cls = classify_smart(txt, kb)
            # IG: tap the reel permalink — comment is visible inline on the reel
            classified_ig.append({
                "comment_id": cmt.get("id", ""),
                "dedup_key": dedup_key("ig-comment", cmt.get("id", "")),
                "media_caption": media_caption,
                "username": cmt.get("username", "?"),
                "text": txt,
                "timestamp": cmt.get("timestamp", ""),
                "reply_url": media_permalink,
                "cls": cls,
                "sentiment": classify_sentiment(txt),
                "urgent": is_urgent(txt),
            })

    print(f"Classified: messenger A/B/NEG = "
          f"{sum(1 for c in classified_messenger if c['cls']['bucket']=='A')}/"
          f"{sum(1 for c in classified_messenger if c['cls']['bucket']=='B')}/"
          f"{sum(1 for c in classified_messenger if c['cls']['bucket']=='NEG')}")

    # 2026-07-03 — LLM-down alert. If the Claude engine was attempted but NEVER
    # succeeded this run (credit-low 400 or any hard error), SMS Lauren — at most
    # once per day (docs/meta/llm_alert_state.json) so it does not spam every 3h.
    try:
        _h = _LLM_HEALTH
        _alert_path = Path(__file__).resolve().parent.parent / "docs/meta/llm_alert_state.json"
        if _h["attempts"] > 0 and _h["ok"] == 0:
            import datetime as _dtm
            _today = _dtm.datetime.now(_dtm.timezone.utc).strftime("%Y-%m-%d")
            try:
                _st = json.loads(_alert_path.read_text()) if _alert_path.exists() else {}
            except Exception:
                _st = {}
            if _st.get("last_alert_date") != _today:
                if _h["credit_low"]:
                    _body = ("⚠️ מנוע ה-AI של תיבת המסרים כבוי — אזלו הקרדיטים ב-Anthropic.\n"
                             "עובר על מילות מפתח בלבד (פחות חכם).\n"
                             "תיקון: console.anthropic.com ← Plans & Billing")
                else:
                    _body = ("⚠️ מנוע ה-AI של תיבת המסרים נכשל בכל הקריאות.\n"
                             + "שגיאה: " + (_h["last_err"][:120]) + "\n"
                             + "עובר על מילות מפתח בינתיים.")
                try:
                    import lauren_sms as _sms
                    if _sms.LAUREN_PHONE and os.environ.get("SIMPLETEXTING_TOKEN"):
                        _sms.send_sms(_sms.LAUREN_PHONE, _body)
                        print(f"  ✉ LLM-down SMS sent (credit_low={_h['credit_low']})")
                except Exception as _se:
                    print(f"  ⚠ LLM-down SMS failed: {_se}")
                _st["last_alert_date"] = _today
                _st["credit_low"] = _h["credit_low"]
                _st["last_err"] = _h["last_err"]
                try:
                    _alert_path.write_text(json.dumps(_st, ensure_ascii=False, indent=2))
                except Exception:
                    pass
        elif _h["ok"] > 0 and _alert_path.exists():
            try:
                _st = json.loads(_alert_path.read_text())
                if _st.get("last_alert_date"):
                    _st["last_alert_date"] = ""
                    _alert_path.write_text(json.dumps(_st, ensure_ascii=False, indent=2))
            except Exception:
                pass
        print(f"  LLM health: {_h['ok']}/{_h['attempts']} ok, credit_low={_h['credit_low']}")
    except Exception as _he:
        print(f"  ⚠ LLM health check failed (non-fatal): {_he}")

    # 2026-05-14 — urgent message alerts. Scans all classified items and
    # SMSes Lauren immediately on any urgent flag (not deduplicated against
    # handled — if the URGENT flag fires on a new message, Lauren needs to know).
    urgent_items = []
    for m in classified_messenger:
        if m.get("urgent") and not handled.get(m.get("dedup_key", ""), {}).get("urgent_smsed"):
            urgent_items.append(("DM", m.get("name","?"), m.get("msg","")))
    for c in classified_fb:
        if c.get("urgent") and not handled.get(c.get("dedup_key", ""), {}).get("urgent_smsed"):
            urgent_items.append(("FB comment", c.get("from","?"), c.get("text","")))
    for c in classified_ig:
        if c.get("urgent") and not handled.get(c.get("dedup_key", ""), {}).get("urgent_smsed"):
            urgent_items.append(("IG comment", c.get("username","?"), c.get("text","")))

    if urgent_items:
        print(f"  🚨 {len(urgent_items)} urgent messages — sending SMS alert to Lauren")
        try:
            import lauren_sms as _sms
            for source, sender, txt in urgent_items[:5]:  # cap at 5 to avoid SMS spam
                short = (txt or "")[:120]
                body = (
                    f"🚨 URGENT INBOX — {source}\n"
                    f"מ-{sender}:\n"
                    f"{short}\n"
                    f"\n"
                    f"👉 dashboard.themakeupblowout.com/meta/"
                )
                if _sms.LAUREN_PHONE and os.environ.get("SIMPLETEXTING_TOKEN"):
                    try:
                        _sms.send_sms(_sms.LAUREN_PHONE, body)
                        print(f"  ✓ urgent SMS sent: {source} / {sender}")
                    except Exception as e:
                        print(f"  ⚠ urgent SMS failed: {e}")
            # Mark items as urgent_smsed so we don't re-SMS Lauren on each daily run
            for source, sender, txt in urgent_items:
                # Find the matching item back and mark it
                for items, channel in [(classified_messenger, "messenger"),
                                        (classified_fb, "fb-comment"),
                                        (classified_ig, "ig-comment")]:
                    for it in items:
                        if it.get("urgent") and (it.get("name") == sender or it.get("from") == sender or it.get("username") == sender):
                            k = it.get("dedup_key", "")
                            if k:
                                h = handled.get(k, {})
                                h["urgent_smsed"] = True
                                handled[k] = h
        except Exception as e:
            print(f"  ⚠ urgent alert handling failed: {e}")

    # Render preview
    html_out = render_preview(snap, classified_messenger, classified_fb, classified_ig)
    out_dir = repo / "docs/meta/inbox-api-preview"
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "index.html").write_text(html_out, encoding="utf-8")

    # Also save the JSON (machine-readable)
    (out_dir / "data.json").write_text(json.dumps({
        "fetched_at": snap["fetched_at"],
        "messenger": classified_messenger,
        "fb_comments": classified_fb,
        "ig_comments": classified_ig,
        "errors": snap.get("errors", []),
    }, indent=2, ensure_ascii=False), encoding="utf-8")

    print(f"Wrote {out_dir/'index.html'} ({(out_dir/'index.html').stat().st_size} bytes)")
    print(f"Wrote {out_dir/'data.json'}")

    # === Phase 2 — supervised single-reply test ===
    if args.reply_test:
        target = args.reply_test
        match = next((m for m in classified_messenger if m.get("conv_id") == target), None)
        if not match:
            print(f"  ❌ no Bucket A draft found for conv_id={target}")
        elif match["cls"]["bucket"] != "A":
            print(f"  ❌ conv_id={target} is in Bucket {match['cls']['bucket']} — only Bucket A items can be auto-replied")
        elif not match["cls"].get("reply"):
            print(f"  ❌ conv_id={target} has no reply text")
        else:
            from lauren_meta import reply_to_messenger
            # Find the recipient PSID (the customer participant)
            try:
                page_id = get_fb_page_id()
                # We need to fetch the conversation to get participant IDs
                from lauren_meta import fetch_messenger_conversations
                allc = fetch_messenger_conversations(limit=50, only_unread=False)
                full = next((c for c in allc if c.get("id") == target), None)
                psid = ""
                if full:
                    for pp in full.get("participants", {}).get("data", []):
                        if pp.get("id") and pp.get("id") != page_id:
                            psid = pp["id"]
                            break
                if not psid:
                    print(f"  ❌ couldn't resolve recipient PSID for {target}")
                else:
                    print(f"  📤 sending live reply to {match['name']} (PSID={psid})...")
                    print(f"      text: {match['cls']['reply']}")
                    r = reply_to_messenger(psid, match["cls"]["reply"], dry_run=False)
                    print(f"  ✓ sent: {r}")
                    # Mark handled so the next run skips it
                    handled[dedup_key("messenger", target)] = {
                        "handled": True,
                        "handledAt": _now_iso(),
                        "method": "phase2-reply-test",
                    }
                    # Persist (best-effort — the workflow will commit via the deploy step)
                    handled_path = Path(__file__).resolve().parent.parent / "docs/meta/handled.json"
                    handled_path.write_text(json.dumps(handled, indent=2, ensure_ascii=False), encoding="utf-8")
                    print(f"  ✓ marked handled in docs/meta/handled.json")
            except Exception as e:
                print(f"  ❌ reply failed: {e}")

    # === Phase 2 — full Bucket A auto-reply (live) ===
    if args.reply_bucket_a:
        from lauren_meta import reply_to_messenger, fetch_messenger_conversations
        sent = 0; failed = 0
        page_id = get_fb_page_id()
        allc = fetch_messenger_conversations(limit=50, only_unread=False)
        psid_by_conv = {}
        for c in allc:
            for pp in c.get("participants", {}).get("data", []):
                if pp.get("id") and pp.get("id") != page_id:
                    psid_by_conv[c.get("id")] = pp["id"]
                    break
        skipped_24h = 0
        from datetime import datetime, timezone, timedelta
        now_utc = datetime.now(timezone.utc)
        for m in classified_messenger:
            if m["cls"]["bucket"] != "A" or not m["cls"].get("reply"):
                continue
            psid = psid_by_conv.get(m["conv_id"])
            if not psid:
                failed += 1; continue
            # Check 24-hour window — Meta RESPONSE messages only allowed within 24h of customer's last message
            updated_time = m.get("updated_time", "")
            if updated_time:
                try:
                    msg_time = datetime.fromisoformat(updated_time.replace("Z", "+00:00"))
                    age_hours = (now_utc - msg_time).total_seconds() / 3600
                    if age_hours > 24:
                        # Re-bucket as B (needs manual reply via Meta UI)
                        m["cls"]["bucket"] = "B"
                        m["cls"]["reason"] = f"older than 24h ({age_hours:.0f}h) — manual reply required"
                        skipped_24h += 1
                        continue
                except Exception:
                    pass
            try:
                reply_to_messenger(psid, m["cls"]["reply"], dry_run=False)
                handled[dedup_key("messenger", m["conv_id"])] = {
                    "handled": True, "handledAt": _now_iso(), "method": "phase2-auto",
                }
                sent += 1
            except Exception as e:
                print(f"  ❌ {m['name']}: {e}")
                failed += 1
        print(f"  Phase 2 auto-reply: sent={sent} failed={failed} skipped_24h={skipped_24h}")
        handled_path = Path(__file__).resolve().parent.parent / "docs/meta/handled.json"
        handled_path.write_text(json.dumps(handled, indent=2, ensure_ascii=False), encoding="utf-8")

    # === Phase 2b — auto-reply to FB + IG comments (Bucket A) ===
    # 2026-05-14 PM — Lauren's directive: extend auto-reply from DMs to public
    # post comments. Same Bucket-A safety: only KB-answerable, never NEG.
    if args.reply_bucket_a:
        from lauren_meta import reply_to_comment
        c_sent = 0; c_failed = 0
        all_comments = (
            [("fb_comment", c) for c in classified_fb] +
            [("ig_comment", c) for c in classified_ig]
        )
        for channel, c in all_comments:
            if c["cls"]["bucket"] != "A" or not c["cls"].get("reply"):
                continue
            comment_id = c.get("comment_id") or c.get("id")
            if not comment_id:
                c_failed += 1; continue
            key = c.get("dedup_key") or dedup_key(channel, comment_id)
            if handled.get(key, {}).get("handled"):
                continue
            try:
                # 2026-07-02 FIX — IG replies must go to /replies (channel-aware)
                reply_to_comment(comment_id, c["cls"]["reply"], dry_run=False,
                                 ig=(channel == "ig_comment"))
                handled[key] = {
                    "handled": True, "handledAt": _now_iso(),
                    "method": f"phase2b-{channel}-auto",
                }
                c_sent += 1
            except Exception as e:
                print(f"  ❌ {channel} {comment_id}: {e}")
                c_failed += 1
        print(f"  Phase 2b comment auto-reply: sent={c_sent} failed={c_failed}")
        # 2026-07-02 — fail-loud (IRON RULE #3 spirit): the IG /comments-endpoint
        # bug failed EVERY reply for weeks with zero alerts because the run
        # itself "succeeded". 3+ failures in one run → SMS Lauren.
        if c_failed >= 3:
            try:
                from lauren_sms import send_sms
                send_sms(os.environ.get("LAUREN_PHONE", "4243547625"),
                         f"@meta ⚠ {c_failed} תשובות אוטומטיות לתגובות נכשלו בריצה הזו "
                         f"(נשלחו {c_sent}). צריך לבדוק את הלוג.\n"
                         "https://github.com/laurenlev10/lauren-agent-hub-data/actions")
            except Exception as _sms_e:
                print(f"  ⚠ failure-SMS failed (non-fatal): {_sms_e}")
        if c_sent:
            handled_path.write_text(json.dumps(handled, indent=2, ensure_ascii=False), encoding="utf-8")

    # === Phase 2c — Instagram DIRECT MESSAGES (2026-07-03) ===
    # IG DMs are reachable WITHOUT App Review via the Instagram-login token
    # (IG_LOGIN_TOKEN). We POLL conversations each run (no webhook needed),
    # classify the latest customer message with the SAME Claude engine, auto-reply
    # Bucket A, and route B/NEG into the same pending queue Lauren reviews.
    try:
        import lauren_ig_dm as _igdm
        handled_path = Path(__file__).resolve().parent.parent / "docs/meta/handled.json"
        if _igdm.get_token():
            convs = _igdm.fetch_conversations(limit=25)
            print(f"  IG DMs: {len(convs)} conversations")
            d_a = d_b = d_neg = d_skip = d_sent = d_sendfail = 0
            # load pending (merge-safe) to append IG-DM B/NEG items
            _pp = Path(__file__).resolve().parent.parent / "docs/meta/inbox-api-preview/pending.json"
            try:
                _pend = json.loads(_pp.read_text(encoding="utf-8"))
            except Exception:
                _pend = {"_updated_at": None, "items": {}}
            _pitems = _pend.get("items", {})
            for cv in convs:
                cid = cv.get("id", "")
                parts = (cv.get("participants") or {}).get("data", [])
                cust = next((pp for pp in parts if pp.get("id") != _igdm.ME_ID), {})
                cust_user = cust.get("username", "?"); cust_id = cust.get("id", "")
                msgs = _igdm.fetch_messages(cid, limit=8)
                last_cust, answered = _igdm.latest_customer_message(msgs)
                if not last_cust or answered:
                    continue
                msg_id = last_cust.get("id") or (cid + ":" + (last_cust.get("created_time") or ""))
                dkey = dedup_key("ig-dm", msg_id)
                if handled.get(dkey, {}).get("handled"):
                    continue
                txt = (last_cust.get("message") or "").strip()
                if not txt:
                    continue
                kb["_seed"] = msg_id; kb["_post_context"] = None
                cls = classify_smart(txt, kb)
                bucket = cls.get("bucket")
                if bucket == "A" and cls.get("reply"):
                    d_a += 1
                    if args.reply_bucket_a and cust_id:
                        try:
                            _igdm.send_reply(cust_id, cls["reply"], dry_run=False)
                            handled[dkey] = {"handled": True, "handledAt": _now_iso(),
                                             "method": "phase2c-igdm-auto"}
                            d_sent += 1
                            continue
                        except Exception as _e:
                            # e.g. outside the 24h messaging window -> hand to Lauren
                            print(f"  ⚠ ig-dm send {cust_user}: {str(_e)[:90]}")
                            d_sendfail += 1
                            _pitems[dkey] = {"channel": "IG DM", "who": "@"+cust_user,
                                "what": txt, "url": f"https://instagram.com/{cust_user}",
                                "bucket": "B", "dedup_key": dkey, "reply_kind": "ig_dm",
                                "target_id": cust_id, "received": last_cust.get("created_time",""),
                                "event_chip": cls.get("event_chip"), "draft": cls.get("reply") or "",
                                "first_seen": _pitems.get(dkey, {}).get("first_seen") or _now_iso()}
                elif bucket in ("B", "NEG"):
                    d_b += (bucket == "B"); d_neg += (bucket == "NEG")
                    _pitems[dkey] = {"channel": "IG DM", "who": "@"+cust_user,
                        "what": txt, "url": f"https://instagram.com/{cust_user}",
                        "bucket": bucket, "dedup_key": dkey, "reply_kind": "ig_dm",
                        "target_id": cust_id, "received": last_cust.get("created_time",""),
                        "event_chip": cls.get("event_chip"), "draft": cls.get("reply") or "",
                        "first_seen": _pitems.get(dkey, {}).get("first_seen") or _now_iso()}
                else:
                    d_skip += 1
            _pend["items"] = _pitems; _pend["_updated_at"] = _now_iso()
            _pp.write_text(json.dumps(_pend, ensure_ascii=False, indent=1), encoding="utf-8")
            if d_sent or d_a or d_b or d_neg:
                handled_path.write_text(json.dumps(handled, indent=2, ensure_ascii=False), encoding="utf-8")
            print(f"  Phase 2c IG DMs: A={d_a} (sent={d_sent}, sendfail={d_sendfail}) B={d_b} NEG={d_neg} SKIP={d_skip}")
        else:
            print("  IG DMs: no IG_LOGIN_TOKEN set — skipping")
    except Exception as _ige:
        print(f"  ⚠ Phase 2c IG DM handling failed (non-fatal): {_ige}")

    # === Bulk mark all classified items as handled (manual cleanup) ===
    if args.bulk_mark_handled:
        marked = 0
        for m in classified_messenger:
            key = dedup_key("messenger", m["conv_id"])
            if not handled.get(key, {}).get("handled"):
                handled[key] = {"handled": True, "handledAt": _now_iso(),
                                "method": "bulk-mark-manual-cleanup"}
                marked += 1
        for c in classified_fb:
            key = c.get("dedup_key", "")
            if key and not handled.get(key, {}).get("handled"):
                handled[key] = {"handled": True, "handledAt": _now_iso(),
                                "method": "bulk-mark-manual-cleanup"}
                marked += 1
        for c in classified_ig:
            key = c.get("dedup_key", "")
            if key and not handled.get(key, {}).get("handled"):
                handled[key] = {"handled": True, "handledAt": _now_iso(),
                                "method": "bulk-mark-manual-cleanup"}
                marked += 1
        handled_path = Path(__file__).resolve().parent.parent / "docs/meta/handled.json"
        handled_path.write_text(json.dumps(handled, indent=2, ensure_ascii=False), encoding="utf-8")
        print(f"  ✓ bulk-mark-handled: marked {marked} items as handled")

    if args.send_sms:
        try:
            sys.path.insert(0, str(Path(__file__).resolve().parent))
            from lauren_sms import send_sms
            need_lauren = (
                sum(1 for c in classified_messenger if c["cls"]["bucket"] in ("B", "NEG")) +
                sum(1 for c in classified_fb        if c["cls"]["bucket"] in ("B", "NEG")) +
                sum(1 for c in classified_ig        if c["cls"]["bucket"] in ("B", "NEG"))
            )
            has_neg = (
                any(c["cls"]["bucket"] == "NEG" for c in classified_messenger) or
                any(c["cls"]["bucket"] == "NEG" for c in classified_fb) or
                any(c["cls"]["bucket"] == "NEG" for c in classified_ig)
            )
            url = "https://dashboard.themakeupblowout.com/meta/inbox-api-preview/"
            if has_neg:
                t3 = f"@meta ⚠ יש תגובה שלילית! {need_lauren} פריטים דורשים אותך.\n{url}"
            elif need_lauren > 0:
                t3 = f"@meta ✓ {need_lauren} פריטים דורשים אותך — לחיצה אחת לכל אחד.\n{url}"
            else:
                t3 = f"@meta ✓ הכל שקט. אין מה לטפל.\n{url}"
            phone = os.environ.get("LAUREN_PHONE", "4243547625")
            r = send_sms(phone, t3)
            print(f"  SMS sent: id={r.get('id')}")
        except Exception as e:
            print(f"  ⚠ SMS failed (non-fatal): {e}")


if __name__ == "__main__":
    main()
